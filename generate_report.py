from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── PAGE SETUP ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin   = Cm(3.75)   # 37.5mm (spec: 35–40mm)
section.right_margin  = Cm(2.25)   # 22.5mm (spec: 20–25mm)
section.top_margin    = Cm(3.25)   # 32.5mm (spec: 30–35mm)
section.bottom_margin = Cm(2.75)   # 27.5mm (spec: 25–30mm)

# ── HELPERS ─────────────────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=False,
             italic=False, space_before=0, space_after=6, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.keep_with_next = keep_with_next
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    """parts = list of (text, bold, italic, size)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for text, bold, italic, size in parts:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic)
    return p

def body(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Cm(2.0)   # 20mm offset from left margin (spec requirement)
    r = p.add_run(text)
    set_font(r, 12)
    return p

def double_body(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Double-spaced body paragraph — used for Abstract and Bonafide Certificate (spec requirement)."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(text)
    set_font(r, 12)
    return p

def bullet(text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size)
    return p

def numbered(text, size=12):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size)
    return p

def chapter_heading(num, title):
    doc.add_page_break()
    # 50mm from top: top_margin=32.5mm, so add ~50pt space_before to reach 50mm
    add_para(f"CHAPTER {num}", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True,
             space_before=50, space_after=0)
    add_para(title.upper(), WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True,
             space_before=0, space_after=24)

def appendix_heading(num, title):
    doc.add_page_break()
    add_para(f"APPENDIX {num}", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True,
             space_before=50, space_after=0)
    add_para(title.upper(), WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True,
             space_before=0, space_after=24)

def section_heading(number, title):
    add_para(f"{number} {title.upper()}", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True,
             space_before=12, space_after=6)

def subsection_heading(number, title):
    add_para(f"{number} {title}", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True,
             space_before=8, space_after=4)

def setup_page_numbers(sec, fmt='decimal', start=1, suppress_first=False):
    """Add right-aligned page numbers to section header. fmt: 'lowerRoman' or 'decimal'."""
    sectPr = sec._sectPr
    # Remove any existing pgNumType
    for existing in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(existing)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

    if suppress_first:
        # Enable different first-page header
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)

    header = sec.header
    header.is_linked_to_previous = False

    # Clear existing content
    for p in header.paragraphs:
        p.clear()

    # Right-aligned page number field
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    set_font(run, 12)

    if suppress_first:
        # First-page header stays blank (title page shows no number)
        # w:titlePg is already set above; the first-page header will be blank by default
        pass

def img_placeholder(fig_num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'D9D9D9')
    p._p.get_or_add_pPr().append(shading)
    r = p.add_run(f"\n\n[INSERT SCREENSHOT / DIAGRAM HERE]\n{fig_num}: {caption}\n\n")
    set_font(r, 11, italic=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    rc = cap.add_run(f"Figure {fig_num}: {caption}")
    set_font(rc, 11, italic=True)

def add_table(headers, rows, caption_num, caption_text):
    # Caption ABOVE table (spec requirement: table titles appear above tables)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after  = Pt(4)
    rc = cap.add_run(f"Table {caption_num}: {caption_text}")
    set_font(rc, 11, bold=True)
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_font(r, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'D9D9D9')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(cell_text))
            set_font(r, 11)
    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4): add_para("")
add_para("PHISHGUARDAI", WD_ALIGN_PARAGRAPH.CENTER, 16, bold=True)
add_para("AI-POWERED REAL-TIME PHISHING DETECTION SYSTEM", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True)
add_para("", space_after=6)
add_para("By", WD_ALIGN_PARAGRAPH.CENTER, 12)
add_para("")
add_para("[YOUR NAME]", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True)
add_para("Roll No. [ROLL NUMBER]  Reg. No. [REGISTER NUMBER]", WD_ALIGN_PARAGRAPH.CENTER, 12)
add_para("")
add_para("")
add_para("A PROJECT REPORT", WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True)
add_para("Submitted to the", WD_ALIGN_PARAGRAPH.CENTER, 12)
add_para("")
add_para("FACULTY OF INFORMATION AND COMMUNICATION ENGINEERING", WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True)
add_para("")
add_para("in partial fulfillment for the award of the degree", WD_ALIGN_PARAGRAPH.CENTER, 12, italic=True)
add_para("Of", WD_ALIGN_PARAGRAPH.CENTER, 12, italic=True)
add_para("")
add_para("MASTER OF COMPUTER APPLICATIONS", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True)
add_para("")
add_para("")
add_para("[UNIVERSITY LOGO PLACEHOLDER]", WD_ALIGN_PARAGRAPH.CENTER, 11, italic=True)
add_para("")
add_para("")
add_para("CENTRE FOR DISTANCE EDUCATION", WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True)
add_para("")
add_para("ANNA UNIVERSITY", WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True)
add_para("")
add_para("CHENNAI 600 025", WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True)
add_para("")
add_para("April 2026", WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True)

# Preliminary pages: Roman numerals (i on title page but suppressed, ii onward shown)
setup_page_numbers(section, fmt='lowerRoman', start=1, suppress_first=True)

# ════════════════════════════════════════════════════════════════════════════
# BONAFIDE CERTIFICATE
# ════════════════════════════════════════════════════════════════════════════
page_break()
add_para("BONAFIDE CERTIFICATE", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=18)
body("Certified that the project report titled ")
# Rebuild with mixed formatting
doc.paragraphs[-1]._p.getparent().remove(doc.paragraphs[-1]._p)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE   # spec: double spacing for Bonafide Certificate
r1 = p.add_run("Certified that the project report titled ")
set_font(r1, 12)
r2 = p.add_run('"PhishGuardAI – AI-Powered Real-Time Phishing Detection System"')
set_font(r2, 12, bold=True)
r3 = p.add_run(" is the Bonafide work of Mr./Ms. ")
set_font(r3, 12)
r4 = p.add_run("[YOUR NAME]")
set_font(r4, 12, bold=True)
r5 = p.add_run(" who carried out the work under my supervision. Certified further that to the best of my knowledge the work reported herein does not form part of any other project report or dissertation on the basis of which a degree or award was conferred on an earlier occasion on this or any other candidate.")
set_font(r5, 12)

add_para("")
add_para("Date : ___________________", WD_ALIGN_PARAGRAPH.LEFT, 12, space_before=12)
for _ in range(6): add_para("")

sig_table = doc.add_table(rows=3, cols=2)
sig_table.style = 'Table Grid'
sig_table.style = doc.styles['Normal Table']
cells = sig_table.rows
for r_idx, (left, right) in enumerate([
    ("Signature of Student", "Signature of Guide"),
    ("[YOUR NAME]", "[GUIDE NAME]"),
    ("Roll No: [ROLL NUMBER]\nRegister No: [REGISTER NUMBER]",
     "[Designation]\n[Department]\n[College]\nAnna University, Chennai – 25")
]):
    lc = sig_table.rows[r_idx].cells[0]
    rc = sig_table.rows[r_idx].cells[1]
    lc.text = ""; rc.text = ""
    rl = lc.paragraphs[0].add_run(left)
    rr = rc.paragraphs[0].add_run(right)
    bold_flag = r_idx == 1
    set_font(rl, 12, bold=bold_flag)
    set_font(rr, 12, bold=bold_flag)

# ════════════════════════════════════════════════════════════════════════════
# CERTIFICATE OF VIVA-VOCE
# ════════════════════════════════════════════════════════════════════════════
page_break()
add_para("CERTIFICATE OF VIVA – VOCE EXAMINATION", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=18)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE   # spec: double spacing for certificates
r1 = p.add_run("This is to certify that Mr./Ms. ")
set_font(r1, 12)
r2 = p.add_run("[YOUR NAME]")
set_font(r2, 12, bold=True)
r3 = p.add_run(" (Roll No: [ROLL NUMBER], Register No: [REGISTER NUMBER]) has been subjected to Viva – voce Examination on ___________ at the Study Centre: Centre for Distance Education, Anna University.")
set_font(r3, 12)
for _ in range(5): add_para("")

viva_table = doc.add_table(rows=4, cols=2)
for r_idx, (left, right) in enumerate([
    ("Internal Examiner", "External Examiner"),
    ("Name :", "Name :"),
    ("Designation :", "Designation :"),
    ("Address :", "Address :")
]):
    lc = viva_table.rows[r_idx].cells[0]
    rc = viva_table.rows[r_idx].cells[1]
    lc.text = ""; rc.text = ""
    rl = lc.paragraphs[0].add_run(left)
    rr = rc.paragraphs[0].add_run(right)
    set_font(rl, 12, bold=(r_idx == 0))
    set_font(rr, 12, bold=(r_idx == 0))

for _ in range(3): add_para("")
add_para("COORDINATOR STUDY CENTER", WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True, space_before=12)
add_para("Name         :", WD_ALIGN_PARAGRAPH.CENTER, 12, space_before=8)
add_para("Designation :", WD_ALIGN_PARAGRAPH.CENTER, 12)
add_para("Address      :", WD_ALIGN_PARAGRAPH.CENTER, 12)

# ════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ════════════════════════════════════════════════════════════════════════════
page_break()
add_para("ACKNOWLEDGEMENT", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=18)

ack_paras = [
    "I extend my sincere gratitude to the faculty at Anna University, Chennai for their invaluable support and guidance throughout this project. I would like to express my deep appreciation to my course instructor at the Centre for Distance Education, Anna University, Chennai – 25, whose encouragement and insightful feedback were instrumental throughout the course of this work.",
    "I am especially grateful to my project guide for their continuous support and highly valuable technical suggestions. Their expertise in Artificial Intelligence, Cybersecurity, Natural Language Processing, and Web Technologies significantly enriched the quality and depth of this project. Every discussion and review session proved to be immensely helpful in shaping the system architecture and implementation strategy.",
    "I also extend my sincere thanks to Anna University, Centre for Distance Education, for providing a stimulating academic environment and access to resources that made this research possible. The knowledge and skills acquired through this program have been invaluable.",
    "My heartfelt thanks go to my classmates and peers who shared ideas, participated in testing sessions, and provided constructive feedback that helped improve the system iteratively. Their enthusiasm for technology and collaborative spirit made this journey enjoyable and productive.",
    "Finally, I want to express my deepest gratitude to my parents and family members for their constant support, patience, and encouragement throughout this journey. They have been my backbone on any given day, and this achievement is as much theirs as it is mine.",
]
for t in ack_paras:
    body(t)
    add_para("")
add_para("([YOUR NAME])", WD_ALIGN_PARAGRAPH.RIGHT, 12, bold=True, space_before=18)

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
page_break()
add_para("TABLE OF CONTENTS", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=12)

toc_data = [
    ("Chapter No.", "Title", "Page No."),
    ("1", "INTRODUCTION", ""),
    ("", "1.1 Overview", ""),
    ("", "1.2 Objectives", ""),
    ("", "1.3 Problem Statement", ""),
    ("", "1.4 Literature Survey", ""),
    ("", "1.5 Existing System", ""),
    ("", "1.6 Need for the System", ""),
    ("", "1.7 Scope of the Project", ""),
    ("2", "REQUIREMENTS SPECIFICATION", ""),
    ("", "2.1 Overview Description", ""),
    ("", "2.2 Key Features", ""),
    ("", "2.3 Hardware and Software Requirements", ""),
    ("", "2.4 Use Case Diagram", ""),
    ("", "2.5 Data Flow Diagram", ""),
    ("3", "SYSTEM DESIGN", ""),
    ("", "3.1 Proposed System", ""),
    ("", "3.2 System Architecture", ""),
    ("", "3.3 Repository Structure", ""),
    ("", "3.4 API Design", ""),
    ("", "3.5 Database Schema", ""),
    ("", "3.6 List of Modules", ""),
    ("", "3.7 Overall System Design", ""),
    ("", "3.8 System Testing", ""),
    ("", "3.9 Test Cases", ""),
    ("4", "IMPLEMENTATION AND ANALYSIS OF RESULTS", ""),
    ("", "4.1 Implementation Environment", ""),
    ("", "4.2 Technology Stack", ""),
    ("", "4.3 Modules Completed", ""),
    ("", "4.4 Deviations and Justifications", ""),
    ("", "4.5 Project Roadmap and Status", ""),
    ("5", "CONCLUSION AND FUTURE WORK", ""),
    ("", "5.1 Limitations", ""),
    ("", "5.2 Future Enhancement", ""),
    ("", "5.3 Conclusion", ""),
    ("", "References", ""),
    ("", "Appendices", ""),
]
toc_table = doc.add_table(rows=len(toc_data), cols=3)
toc_table.style = 'Table Grid'
for ri, (col1, col2, col3) in enumerate(toc_data):
    row = toc_table.rows[ri]
    for ci, txt in enumerate([col1, col2, col3]):
        row.cells[ci].text = ""
        is_bold = (ri == 0) or (col1.isdigit() and ci == 0) or (ci == 1 and col1.isdigit())
        r = row.cells[ci].paragraphs[0].add_run(txt)
        set_font(r, 11, bold=is_bold)
        if ci == 2:
            row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ci == 0:
            row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ════════════════════════════════════════════════════════════════════════════
page_break()
add_para("LIST OF TABLES", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=12)
lot_data = [
    ("Table No.", "Title of Table", "Page No."),
    ("2.1", "Hardware Requirements", ""),
    ("2.2", "Software Requirements", ""),
    ("3.1", "API Endpoint Summary", ""),
    ("3.2", "Database Entity Relationships", ""),
    ("3.3", "List of Modules", ""),
    ("3.4", "Test Cases – Authentication Module", ""),
    ("3.5", "Test Cases – URL Checking Module", ""),
    ("3.6", "Test Cases – AI Analysis Module", ""),
    ("3.7", "Test Cases – Chrome Extension Module", ""),
    ("3.8", "Test Cases – Admin Dashboard Module", ""),
    ("4.1", "Technology Stack", ""),
    ("4.2", "Modules Completed Summary", ""),
    ("4.3", "Deviations and Justifications", ""),
    ("4.4", "Project Roadmap and Phase Status", ""),
]
lot_table = doc.add_table(rows=len(lot_data), cols=3)
lot_table.style = 'Table Grid'
for ri, row_data in enumerate(lot_data):
    for ci, txt in enumerate(row_data):
        lot_table.rows[ri].cells[ci].text = ""
        r = lot_table.rows[ri].cells[ci].paragraphs[0].add_run(txt)
        set_font(r, 11, bold=(ri == 0))
        if ci in (0, 2):
            lot_table.rows[ri].cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ════════════════════════════════════════════════════════════════════════════
page_break()
add_para("LIST OF FIGURES", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=12)
lof_data = [
    ("Figure No.", "Title of Figure", "Page No."),
    ("2.1", "Use Case Diagram", ""),
    ("2.2", "Data Flow Diagram – Level 0 (Context Diagram)", ""),
    ("2.3", "Data Flow Diagram – Level 1 (Detailed)", ""),
    ("3.1", "System Architecture Diagram", ""),
    ("3.2", "Project Repository Structure", ""),
    ("3.3", "Entity Relationship (ER) Diagram", ""),
    ("3.4", "Prisma Database Schema", ""),
    ("4.1", "User Authentication Flow (Login & Register)", ""),
    ("4.2", "Chrome Extension Popup Interface", ""),
    ("4.3", "URL Scanning Page", ""),
    ("4.4", "Phishing Blocked Page", ""),
    ("4.5", "Warning Page Interface", ""),
    ("4.6", "Admin Dashboard – URL Checks Table", ""),
    ("4.7", "Admin Dashboard – Audit Records Table", ""),
    ("4.8", "LangChain AI Analysis Flow Diagram", ""),
]
lof_table = doc.add_table(rows=len(lof_data), cols=3)
lof_table.style = 'Table Grid'
for ri, row_data in enumerate(lof_data):
    for ci, txt in enumerate(row_data):
        lof_table.rows[ri].cells[ci].text = ""
        r = lof_table.rows[ri].cells[ci].paragraphs[0].add_run(txt)
        set_font(r, 11, bold=(ri == 0))
        if ci in (0, 2):
            lof_table.rows[ri].cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════════════════════════════════
# ABSTRACT (ENGLISH)
# ════════════════════════════════════════════════════════════════════════════
page_break()
add_para("ABSTRACT", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=18)
abstract_paras = [
    "In the modern digital landscape, phishing attacks represent one of the most pervasive and financially devastating cybersecurity threats. With over 3.4 billion phishing attempts occurring daily worldwide and an average organizational loss of $4.65 million per incident, traditional blacklist-based and heuristic detection methods have proven inadequate against the rapidly evolving tactics of modern adversaries.",
    'This project, titled "PhishGuardAI – AI-Powered Real-Time Phishing Detection System," presents an intelligent, browser-integrated security solution that leverages cutting-edge Artificial Intelligence to detect and prevent phishing attacks in real time. The system operates as a Chrome browser extension (Manifest V3) seamlessly integrated with a Next.js 16 backend API and a PostgreSQL database.',
    "The core detection engine integrates Google Gemini Large Language Model through the LangChain framework, enabling multi-turn conversational analysis and tool-calling capabilities. When a user navigates to a URL, the system performs a multi-layered analysis: first checking an in-memory session cache, then querying a shared PostgreSQL database cache for previously analyzed URLs, and finally invoking the AI engine for new URLs. The AI engine examines URL structure, domain patterns, page content, JavaScript behavior, form elements, and link destinations to classify websites.",
    "The system provides three levels of protection: automatic blocking of confirmed phishing sites, configurable warning banners for risky websites (categorized as PIRACY, SCAMMING, RISKY_LINKS, SCAM_PRODUCTS, or RISKY_FILES), and transparent navigation for safe sites. A confidence scoring mechanism (0.0–1.0) guides decision-making, while a comprehensive admin dashboard provides administrators with real-time monitoring, URL analysis history, and detailed audit records.",
    "The solution is built on a full-stack TypeScript architecture using Next.js 16, Prisma ORM, JWT authentication, bcryptjs password hashing, and Radix UI components. The system demonstrates how modern LLM-based intelligence can be deployed as a practical, real-time security tool that adapts to novel attack patterns without requiring manual rule updates or retraining.",
]
for t in abstract_paras:
    double_body(t)   # spec: Abstract must be double-spaced
    add_para("")

# ════════════════════════════════════════════════════════════════════════════
# ABSTRACT (TAMIL)
# ════════════════════════════════════════════════════════════════════════════
page_break()
add_para("சுருக்கம்", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=18)
tamil_paras = [
    "இன்றைய டிஜிட்டல் காலகட்டத்தில், ஃபிஷிங் தாக்குதல்கள் மிகவும் அதிகமான இணைய பாதுகாப்பு அச்சுறுத்தல்களில் ஒன்றாக உள்ளன. தினமும் உலகளவில் 3.4 பில்லியனுக்கும் அதிகமான ஃபிஷிங் முயற்சிகள் நடைபெறுகின்றன, மேலும் ஒவ்வொரு நிறுவனத்திற்கும் சராசரியாக $4.65 மில்லியன் இழப்பு ஏற்படுகிறது.",
    '"PhishGuardAI" என்ற இந்த திட்டம், செயற்கை நுண்ணறிவை பயன்படுத்தி உண்மையான நேரத்தில் ஃபிஷிங் தளங்களை கண்டறிந்து தடுக்கும் ஒரு நவீன Chrome பிரவுசர் நீட்டிப்பாகும். Google Gemini பெரிய மொழி மாதிரி மற்றும் LangChain கட்டமைப்பை கொண்டு URL மற்றும் வலைத்தள உள்ளடக்கத்தை ஆழமாக பகுப்பாய்வு செய்கிறது.',
    "அமைப்பு மூன்று நிலைகளில் பாதுகாப்பை வழங்குகிறது: உறுதிப்படுத்தப்பட்ட ஃபிஷிங் தளங்களை தானாக தடுக்கும், ஆபத்தான தளங்களுக்கு எச்சரிக்கை செய்திகளை காட்டும், மற்றும் பாதுகாப்பான தளங்களுக்கு தடையின்றி வழிசெலுத்தும். நம்பகத்தன்மை மதிப்பெண் (0.0–1.0) முடிவெடுக்க உதவுகிறது.",
    "Next.js 16, PostgreSQL, Prisma ORM, JWT அங்கீகாரம் மற்றும் TypeScript ஆகியவற்றை பயன்படுத்தி கட்டப்பட்ட இந்த அமைப்பு, நவீன AI-சக்தியுள்ள பாதுகாப்பு தீர்வின் நடைமுறை பயன்பாட்டை நிரூபிக்கிறது.",
    "[Note: Tamil abstract placeholder — please verify/complete the Tamil translation with a language expert before final submission.]",
]
for t in tamil_paras:
    double_body(t)   # spec: Abstract must be double-spaced
    add_para("")

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# (New section here: Arabic page numbering restarts at 1 for main text)
# ════════════════════════════════════════════════════════════════════════════
main_section = doc.add_section()
main_section.left_margin   = Cm(3.75)
main_section.right_margin  = Cm(2.25)
main_section.top_margin    = Cm(3.25)
main_section.bottom_margin = Cm(2.75)
setup_page_numbers(main_section, fmt='decimal', start=1)
# Chapter heading without extra page break (section break above already starts new page)
add_para("CHAPTER I", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=0)
add_para("INTRODUCTION", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=0, space_after=24)

section_heading("1.1", "OVERVIEW")
body("In the contemporary digital landscape, phishing attacks have emerged as one of the most significant and persistent cybersecurity threats facing individuals, organizations, and critical infrastructure worldwide. Unlike many other forms of cyberattacks that exploit technical vulnerabilities in software systems, phishing specifically targets human psychology — manipulating users into voluntarily divulging sensitive information such as login credentials, financial details, and personal identification data.")
add_para("")
body("The scale of phishing activity is staggering by any measure. According to recent cybersecurity reports, over 3.4 billion phishing emails are sent daily, and approximately 90% of all data breaches can be traced back to some form of phishing attempt. The financial implications are equally alarming, with the average cost of a successful phishing attack reaching $4.65 million per incident for enterprises. In 2023 alone, the FBI's Internet Crime Complaint Center (IC3) reported phishing as the most common cybercrime type, with losses exceeding $52 million directly attributed to phishing schemes.")
add_para("")
body("Traditional security mechanisms have demonstrated critical limitations in combating modern phishing campaigns. Blacklist-based systems — the foundation of tools like Google Safe Browsing and Microsoft SmartScreen — can only flag websites that have been previously reported and verified as malicious. This inherently reactive approach creates a dangerous window of vulnerability during which new phishing sites operate freely before being identified. Research indicates that the average phishing site remains active for only 24–48 hours before being reported, meaning the blacklist update cycle is often too slow to protect users from novel attacks.")
add_para("")
body("Heuristic-based detection systems offer an improvement by analyzing URL patterns and structural characteristics to identify suspicious sites. However, these rule-based approaches suffer from high false positive rates that frustrate users, and sophisticated attackers have learned to craft phishing pages that deliberately evade known heuristic patterns. Machine learning approaches, while promising, require substantial training datasets and struggle with generalization to new attack patterns without expensive model retraining.")
add_para("")
body('This project, titled "PhishGuardAI – AI-Powered Real-Time Phishing Detection System," addresses these limitations by leveraging the emergent reasoning capabilities of Large Language Models (LLMs). Specifically, the system integrates Google Gemini LLM through the LangChain framework — a combination that enables genuine semantic understanding of website content, contextual reasoning about threat indicators, and intelligent decision-making that can adapt to novel attack patterns without explicit rule programming or model retraining.')
add_para("")
body("The proposed system operates as a Chrome browser extension that intercepts every URL navigation, performing multi-layered analysis through an in-memory cache, shared database cache, and AI-powered analysis pipeline. The result is a comprehensive security solution that provides real-time protection, shares threat intelligence across users, maintains detailed audit trails for security analysis, and includes an administrative dashboard for system monitoring.")
add_para("")
body("The project is structured as follows: Chapter 2 provides detailed requirements specification including hardware/software requirements, use case diagrams, and data flow diagrams. Chapter 3 presents the complete system design including architecture, API design, database schema, and module descriptions. Chapter 4 covers implementation details, technology stack, completed modules, and deviations. Chapter 5 concludes with limitations, future enhancements, and overall conclusions.")

section_heading("1.2", "OBJECTIVES")
body("The primary objectives of this project are defined as follows, encompassing both the technical and practical goals of the system:")
numbered("To develop a Chrome browser extension (Manifest V3) that automatically intercepts every URL navigation and performs multi-layered phishing analysis without requiring manual user intervention, providing seamless and transparent protection.")
numbered("To integrate Google Gemini Large Language Model through the LangChain framework, enabling intelligent semantic analysis of URLs and website content with tool-calling capabilities for deep inspection when initial URL analysis yields low confidence results.")
numbered("To implement a multi-layered detection mechanism capable of identifying confirmed phishing sites, categorizing risky websites into five distinct warning types (PIRACY, SCAMMING, RISKY_LINKS, SCAM_PRODUCTS, RISKY_FILES), and assigning confidence scores for each assessment.")
numbered("To design and implement a three-level caching architecture — in-memory session cache, shared PostgreSQL database cache, and URL normalization — enabling sub-50ms response times for previously analyzed domains and reducing redundant AI API calls.")
numbered("To develop a comprehensive administrative dashboard providing real-time monitoring of all URL analyses, paginated audit records with filtering capabilities, detection statistics, and system health visibility for security administrators.")
numbered("To implement robust, production-grade security measures including JWT-based stateless authentication, Bcrypt password hashing, role-based access control (USER/ADMIN roles), and fail-open error handling that prioritizes user experience without compromising security integrity.")
numbered("To ensure a seamless user experience with minimal performance overhead, such that phishing protection operates transparently for safe sites while providing clear, informative blocking and warning interfaces for detected threats.")
numbered("To evaluate system performance through comprehensive testing covering authentication flows, URL interception accuracy, AI detection quality, cache effectiveness, and end-to-end integration from URL navigation to block/warn/allow action.")

section_heading("1.3", "PROBLEM STATEMENT")
body("Internet users, employees, students, and digital consumers regularly face the risk of encountering phishing websites during their everyday browsing activities. The current process of protecting against phishing relies primarily on the following sequential mechanisms, each of which has fundamental limitations:")
bullet("Manual blacklist checking: Systems like Google Safe Browsing maintain databases of known malicious URLs. These databases are only updated after sites are reported and verified — a process that typically takes 24 to 72 hours — leaving users unprotected during this critical window.")
bullet("Heuristic rule-based detection: Pattern-matching algorithms analyze URL structure, domain age, and character patterns for suspicious indicators. These systems suffer from high false positive rates (flagging legitimate sites) and are routinely evaded by sophisticated attackers who are familiar with common detection rules.")
bullet("Browser security warnings: Modern browsers display basic warnings for known dangerous sites, but these depend entirely on the same blacklist databases described above. They offer no independent, real-time analysis capability.")
bullet("User awareness training: Organizations invest in training users to recognize phishing attempts. However, phishing sites have become convincingly authentic, with legitimate-looking SSL certificates, professional design, and even copied brand assets, making visual detection unreliable even for trained users.")
add_para("")
body("This manual and reactive process presents several critical problems for the cybersecurity community. It is fundamentally retrospective — protection only activates after victims have already been identified and attacks reported. It is computationally shallow — no existing widely-deployed browser security tool performs genuine semantic analysis of website content to understand intent and behavior. It is isolated — threat intelligence discovered by one user is not rapidly shared to protect other users browsing the same domain. It is binary — sites are either known-bad or treated as unknown/safe, with no nuanced risk categorization for sites that are risky but not confirmed phishing.")
add_para("")
body("Existing commercial solutions that do perform real-time content analysis are typically enterprise-grade products with substantial licensing costs, making them inaccessible to individual users and small organizations. There is therefore a compelling need for an accessible, intelligent, and proactive phishing detection system that can analyze website content in real time, adapt to novel attack patterns through AI reasoning, and share threat intelligence collectively — all while maintaining the ease of use expected from a browser extension.")

section_heading("1.4", "LITERATURE SURVEY")
subsection_heading("1.4.1", "Traditional Phishing Detection and Blacklist-Based Systems")
body("The earliest systematic approaches to phishing detection relied on centralized databases of known malicious URLs. Zhang et al. (2007) proposed CANTINA, one of the first content-based phishing detection systems, which used TF-IDF analysis of webpage content to identify suspicious sites. Their work demonstrated that content analysis could catch sites not yet in blacklists, establishing the theoretical foundation for content-aware detection. The system achieved 97% accuracy on known phishing sites but struggled with zero-day attacks, a limitation that remains relevant to this day.")
add_para("")
body("The APWG (Anti-Phishing Working Group) has documented the evolution of phishing tactics over two decades, showing that modern attacks exploit legitimate cloud hosting, valid SSL certificates, and sophisticated social engineering — all factors that defeat simple URL-pattern analysis. Their reports from 2022-2024 consistently show that the average time from phishing site launch to blacklist inclusion exceeds 24 hours, during which thousands of potential victims may be exposed.")

subsection_heading("1.4.2", "Machine Learning Approaches to Phishing Detection")
body("Sahingoz et al. (2019) conducted an extensive comparative study of seven machine learning algorithms — including Random Forests, Neural Networks, Support Vector Machines, and k-Nearest Neighbours — applied to phishing URL classification. Using a dataset of 73,575 URLs with 30 features extracted from URL structure, domain characteristics, and WHOIS data, their best model (Random Forest) achieved 97.98% accuracy. However, the study highlighted a critical limitation: ML models trained on historical data require periodic retraining to maintain effectiveness against evolving attack patterns, and their performance degrades significantly on novel phishing techniques not represented in the training data.")
add_para("")
body("Mohammad et al. (2015) proposed a hybrid intelligent system combining intelligent rule-based features with an artificial neural network for phishing website detection. Their approach extracted 30 features from URLs and website content and achieved 92.5% accuracy. The work underscored the importance of multi-feature analysis that goes beyond URL structure alone, a principle directly applied in the current project's multi-factor AI analysis prompt design.")

subsection_heading("1.4.3", "Large Language Models for Security Analysis")
body("The emergence of transformer-based Large Language Models has opened transformative possibilities for cybersecurity analysis. Brown et al. (2020) demonstrated in the GPT-3 paper that sufficiently large language models can perform sophisticated reasoning tasks with minimal task-specific training data — a property known as few-shot learning. This capability is directly relevant to phishing detection, where the model can reason about website legitimacy based on general knowledge of phishing tactics rather than requiring a labeled phishing dataset.")
add_para("")
body("Motlagh et al. (2024) conducted a comprehensive survey of LLM applications in cybersecurity, highlighting their potential for malware analysis, vulnerability detection, and social engineering detection. Their work demonstrated that LLMs can identify subtle linguistic and structural patterns in malicious content that rule-based systems miss. The survey particularly emphasized the importance of grounding LLM analysis in retrieved evidence — exactly the approach taken in this project through the fetchWebsiteContent tool — to prevent hallucination and ensure that security decisions are based on actual website content rather than model assumptions.")

subsection_heading("1.4.4", "LangChain Framework and Tool-Augmented LLMs")
body("Chase (2022) introduced LangChain as a framework for building applications powered by language models with tool-calling capabilities. The framework's ReAct (Reasoning + Acting) pattern — where the model reasons about what information it needs, calls a tool to retrieve it, observes the result, and continues reasoning — is particularly valuable for security analysis tasks that require evidence gathering before making a judgment.")
add_para("")
body("The tool-calling paradigm overcomes a key limitation of pure LLM analysis: models cannot autonomously browse the web or inspect live website content. By providing the AI with a structured fetchWebsiteContent tool that extracts page titles, scripts, forms, links, and suspicious patterns from any URL, the system enables genuine evidence-based analysis that mirrors the investigative process of a human security analyst. The LangChain-Google Generative AI integration used in this project leverages Google Gemini's native function-calling API for efficient, structured tool invocation.")

subsection_heading("1.4.5", "Browser Extension Security Architecture")
body("Carlini et al. (2012) analyzed the security architecture of Chrome browser extensions, identifying the key security boundaries between extension components and web page content. Their work on the principle of least privilege in extension design directly informed the Manifest V3 architecture used in this project, where the background service worker has no direct DOM access to web pages, and content scripts operate in isolated worlds separate from page JavaScript.")
add_para("")
body("Jagpal et al. (2015) studied malicious browser extensions and their detection techniques, providing valuable insights into the security considerations for building trustworthy security extensions. Their finding that users are willing to grant broad permissions to security-oriented extensions underscores the responsibility of security extension developers to implement robust privacy protections — a requirement addressed in this project through minimal permission requests and secure JWT-based authentication.")

subsection_heading("1.4.6", "Caching Strategies for Web Security Systems")
body("Weaver et al. (2016) analyzed the performance characteristics of real-time URL reputation systems, demonstrating that caching is not merely a performance optimization but a security necessity. Without effective caching, the latency introduced by real-time AI analysis (which can range from 1-5 seconds per query) would severely degrade user experience to the point where users would disable protection. Their work established that a two-tier caching approach — local session cache combined with a shared persistent cache — provides the optimal balance between freshness of threat intelligence and response performance.")
add_para("")
body("The multi-level caching architecture in this project implements these findings directly: an in-memory cache provides sub-millisecond responses for URLs visited during the current browser session, while the shared PostgreSQL database cache serves as persistent threat intelligence that benefits all users and persists across browser restarts.")

section_heading("1.5", "EXISTING SYSTEM")
body("An analysis of existing systems for phishing detection and web security reveals several significant limitations that the current project aims to address:")
bullet("Google Safe Browsing API: Google's Safe Browsing service is the most widely deployed phishing protection system, embedded in Chrome, Firefox, and Safari. While it provides reasonable protection against known threats, it relies entirely on crowd-sourced reporting and Google's analysis pipeline. New phishing sites typically take 12-48 hours to be added to the database. The API only checks URLs against a known-bad list — it performs no real-time analysis of site content or structure, meaning entirely new phishing techniques that haven't been seen before offer zero protection.")
bullet("Microsoft SmartScreen: Microsoft's SmartScreen filter, integrated into Edge and Windows Defender, similarly depends on reputation databases updated from community reporting and Microsoft's security research. Like Safe Browsing, it lacks real-time content analysis capabilities and cannot reason about site intent from first principles. Its integration with Active Directory provides enterprise policy enforcement, but this feature is inaccessible to individual users and small organizations.")
bullet("Heuristic URL Analyzers (PhishTank, OpenPhish): Services like PhishTank provide community-curated databases of verified phishing URLs and offer API access for developers. However, these are pure blacklists with no semantic analysis capability. A phishing site on a domain not yet in PhishTank's database receives no warning. Additionally, PhishTank's open contribution model has been exploited by attackers who submit legitimate sites to generate false positives.")
bullet("Commercial Endpoint Security (Symantec, CrowdStrike, Palo Alto): Enterprise security products from these vendors do provide real-time URL analysis and content inspection, but they are priced for corporate environments at $30-100+ per user per year. They require agent installation, centralized management infrastructure, and IT administration expertise — making them completely impractical for individual users, students, or small businesses.")
bullet("Browser Privacy Extensions (uBlock Origin, Privacy Badger): These tools focus on ad blocking and tracker prevention rather than phishing detection. While they incidentally block some known malicious domains through community filter lists, they are not designed for phishing detection and miss most phishing sites, which typically use clean domains without tracking infrastructure.")
add_para("")
body("The fundamental gap in all existing systems is the absence of intelligent, real-time semantic analysis — the ability to examine a previously unknown website and reason from its content, structure, and behavior whether it is attempting to deceive users into disclosing sensitive information.")

section_heading("1.6", "NEED FOR THE SYSTEM")
body("The need for an AI-powered phishing detection system is supported by multiple converging trends in cybersecurity, user behavior, and AI capability:")
bullet("Zero-Day Phishing Protection: With the average phishing site active for only 24-48 hours and blacklist updates lagging by the same duration, a substantial proportion of successful phishing attacks occur against sites not yet in any blacklist. An AI system capable of reasoning about site legitimacy from first principles — analyzing domain patterns, content, scripts, and forms — can provide protection during this critical window when traditional systems are blind.")
bullet("Escalating Attack Sophistication: Modern phishing campaigns invest significant resources in appearing legitimate. They purchase domain names visually similar to targeted brands, obtain valid SSL certificates (making HTTPS meaningless as a trust signal), copy brand assets and page layouts precisely, and use legitimate cloud hosting services to avoid IP reputation blocks. Only semantic analysis that actually reads and reasons about page content can reliably detect these sophisticated attacks.")
bullet("Collective Intelligence Through Shared Caching: Individual users browsing the web generate a collective intelligence resource about phishing threats. A shared database cache means that when any user encounters a phishing site and the AI confirms it, subsequent users visiting the same domain receive instant blocking responses without any additional AI computation cost. This shared intelligence model amplifies protection across the entire user base.")
bullet("Accessibility of Protection: Enterprise-grade real-time analysis tools exist but are financially out of reach for the vast majority of internet users. A browser extension with a free-tier backend provides professional-quality protection to individuals, students, and small organizations who cannot afford commercial security platforms.")
bullet("Audit and Accountability: Organizations need visibility into their security posture — which sites were blocked, which triggered warnings, which users attempted to visit suspicious domains, and what actions were taken. The admin dashboard and audit logging capabilities provide this visibility, enabling security teams to identify trends, investigate incidents, and demonstrate compliance.")
bullet("Educational Value of AI Security: Demonstrating that LLM-based reasoning can be deployed as a practical, real-time security tool has significant implications for the broader cybersecurity field. This project serves as a proof-of-concept for AI-augmented security operations that can inspire further research and development in AI-driven threat detection.")

section_heading("1.7", "SCOPE OF THE PROJECT")
body("The scope of PhishGuardAI encompasses the following functional domains and application areas:")

subsection_heading("1.7.1", "Chrome Browser Extension")
bullet("Manifest V3 compliant service worker that intercepts all URL navigations")
bullet("Multi-level caching (in-memory + database) for performance optimization")
bullet("Blocking page for confirmed phishing sites with threat details")
bullet("Warning banner for risky sites with configurable behavior (block or banner)")
bullet("Popup interface for manual URL checking and authentication")
bullet("Settings page for API URL and auto-scan configuration")

subsection_heading("1.7.2", "Backend API Server")
bullet("Next.js 16 API routes providing RESTful endpoints for all extension operations")
bullet("JWT-based stateless authentication with 7-day token expiry")
bullet("User registration and login with email validation and password requirements")
bullet("URL analysis endpoint with database caching and AI integration")
bullet("Admin-only endpoints for URL history and audit record management")

subsection_heading("1.7.3", "AI Analysis Engine")
bullet("LangChain ChatGoogleGenerativeAI integration with Google Gemini model")
bullet("Multi-turn analysis with tool-calling for website content inspection")
bullet("Phishing classification with confidence scoring (0.0–1.0)")
bullet("Five warning categories: PIRACY, SCAMMING, RISKY_LINKS, SCAM_PRODUCTS, RISKY_FILES")
bullet("Severity assessment (low/medium/high) for each warning type")

subsection_heading("1.7.4", "Database Design and Implementation")
bullet("PostgreSQL database with three core models: User, UrlCheck, AuditRecord")
bullet("Prisma ORM for type-safe queries and schema management")
bullet("URL normalization for consistent cache key generation")
bullet("Indexed queries for fast URL lookups")

subsection_heading("1.7.5", "Admin Panel")
bullet("Tab-based dashboard for URL checks and audit records")
bullet("Paginated tables with search, filtering, and sorting capabilities")
bullet("Confidence score display and warning type visualization")
bullet("User email search in audit records")

subsection_heading("1.7.6", "Scope Limitations")
body("The following aspects are explicitly out of scope for the current implementation:")
bullet("Email phishing detection — the system focuses exclusively on web-based URL analysis")
bullet("Support for Firefox, Safari, or Edge browsers — Chrome (Manifest V3) only")
bullet("Offline functionality — internet connectivity is required for AI analysis")
bullet("Enterprise features such as centralized policy management or Active Directory integration")
bullet("Mobile browser extension support")

subsection_heading("1.7.7", "Application Domains")
body("The system is particularly well-suited for the following use cases:")
bullet("Individual internet users seeking AI-powered protection beyond standard browser security features")
bullet("Educational institutions protecting students and staff from credential-harvesting attacks targeting university login portals")
bullet("Small businesses requiring professional-grade phishing protection without enterprise software costs")
bullet("Security researchers studying phishing trends and AI-based detection effectiveness")
bullet("Organizations needing audit trails and visibility into web browsing security events")

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — REQUIREMENTS SPECIFICATION
# ════════════════════════════════════════════════════════════════════════════
chapter_heading("2", "REQUIREMENTS SPECIFICATION")

section_heading("2.1", "OVERVIEW DESCRIPTION")
body("PhishGuardAI is a full-stack cybersecurity application consisting of a Chrome browser extension frontend and a Next.js 16 backend API server with a PostgreSQL database. The system is designed to provide end-to-end protection against phishing websites through intelligent AI analysis, shared threat intelligence caching, and comprehensive administrative visibility.")
add_para("")
body("The overall protection pipeline operates as follows: the user's Chrome browser, with the PhishGuardAI extension installed and authenticated, intercepts every navigation request. The extension background service worker normalizes the URL to its base domain, checks the in-memory session cache, and — if not cached — queries the backend API. The backend checks the PostgreSQL URL cache; if the domain has been previously analyzed, it returns the cached result instantly. For new URLs, the backend invokes the LangChain-Gemini AI analysis pipeline, which may fetch and inspect the website content if the URL pattern analysis alone yields insufficient confidence. The result is stored in the database and returned to the extension, which takes the appropriate action: blocking, warning, or allowing the navigation.")
add_para("")
body("The system supports three distinct user roles: unauthenticated users who can only access the extension login/register interface; authenticated users who can use the full phishing detection functionality; and administrators who have access to the web-based admin dashboard for monitoring and analysis.")

section_heading("2.2", "KEY FEATURES")
subsection_heading("2.2.1", "Real-Time Phishing Detection")
body("The system's primary feature is real-time analysis of every URL a user navigates to. Unlike browser bookmarks or manually submitted URLs, the detection is fully automatic and transparent — users do not need to take any action for protection to be active. The background service worker in the Chrome extension monitors the chrome.tabs.onUpdated event and initiates the analysis pipeline for every new navigation. The fail-open design ensures that if the API is unavailable or returns an error, navigation proceeds normally so that legitimate browsing is never blocked by a system outage.")

subsection_heading("2.2.2", "AI-Powered Semantic Analysis")
body("The detection engine uses Google Gemini LLM through LangChain to perform genuine semantic analysis of URLs and website content. This goes far beyond simple pattern matching — the AI reasons about the intent of the website based on multiple evidence signals. For URL analysis, the model examines domain name patterns for typosquatting (e.g., 'g00gle.com' or 'paypa1.com'), suspicious subdomain structures, and known phishing URL patterns. When URL analysis alone is insufficient, the fetchWebsiteContent tool is invoked to inspect the actual page, extracting form destinations, JavaScript behavior, external link domains, and suspicious obfuscation patterns.")

subsection_heading("2.2.3", "Multi-Layer Caching Architecture")
body("Performance is critical for a security tool — protection that significantly slows down browsing will be disabled by users. PhishGuardAI implements a three-level caching strategy to achieve near-instant responses for the vast majority of requests. The in-memory cache within the extension's service worker stores results for URLs visited during the current browser session, providing sub-millisecond lookup for recently visited domains. The shared PostgreSQL database cache provides persistent storage of all previous analysis results, meaning any URL analyzed by any user benefits all subsequent users visiting the same domain. URL normalization (removing tracking parameters, standardizing protocols) maximizes cache hit rates.")

subsection_heading("2.2.4", "Warning Categorization System")
body("Recognizing that many dangerous websites do not fit the narrow definition of 'phishing' but still pose significant risks to users, the system implements a multi-category warning system. Phishing sites (sites attempting credential theft or financial fraud) trigger hard blocking with a detailed explanation page. Five categories of risky-but-not-phishing sites trigger configurable warnings: PIRACY (sites hosting pirated content or illegal software), SCAMMING (fake services, fraudulent reviews, deceptive practices), RISKY_LINKS (suspicious redirect chains, link farms), SCAM_PRODUCTS (counterfeit goods, fake marketplaces), and RISKY_FILES (malicious file distribution). Each warning carries a severity level (low/medium/high) to help users make informed navigation decisions.")

subsection_heading("2.2.5", "Admin Dashboard")
body("A web-based administrative interface built with Next.js and Radix UI provides security administrators with comprehensive visibility into system activity. The URL Checks tab displays a paginated, searchable, and filterable table of all analyzed URLs, showing detection results, confidence scores, warning categories, and analysis timestamps. The Audit Records tab provides a detailed log of user interactions with potentially harmful or warned sites, including the user email, URL visited, action taken (BLOCKED/ALLOWED/WARNING_SHOWN), IP address, user agent, and timestamp. Date range filtering and status filtering enable efficient investigation of specific incidents or time periods.")

subsection_heading("2.2.6", "Secure Authentication and Access Control")
body("The authentication system implements industry-standard security practices throughout. User passwords are hashed using bcryptjs with a configurable work factor (default: 10 salt rounds) before storage, ensuring that database breaches cannot expose plaintext credentials. JWT tokens are issued upon successful login with a 7-day expiry, enabling stateless authentication that scales without server-side session storage. Role-based access control with USER and ADMIN roles ensures that the admin dashboard is accessible only to designated administrators. The API middleware validates JWT tokens on every protected request, additionally confirming that the user record still exists in the database to handle account deletion scenarios.")

section_heading("2.3", "HARDWARE AND SOFTWARE REQUIREMENTS")
subsection_heading("2.3.1", "Hardware Requirements")
hw_data = [
    ("Processor", "Intel Core i3 / AMD Ryzen 3 (2.0 GHz, Dual Core)", "Intel Core i5 / AMD Ryzen 5 (3.0 GHz, Quad Core) or higher"),
    ("RAM", "4 GB DDR4", "8 GB DDR4 or higher"),
    ("Storage", "20 GB HDD (for OS, Node.js, DB)", "50 GB SSD (for faster I/O)"),
    ("Network", "Broadband Internet (for Gemini API)", "High-speed broadband (10 Mbps+)"),
    ("Operating System", "Windows 10 / Ubuntu 20.04 LTS", "Ubuntu 22.04 LTS (for production)"),
    ("Browser", "Google Chrome 90+", "Google Chrome 120+ (latest stable)"),
]
add_table(["Component", "Minimum Specification", "Recommended Specification"], hw_data, "2.1", "Hardware Requirements")

subsection_heading("2.3.2", "Software Requirements")
sw_data = [
    ("Runtime", "Node.js", "v20 LTS"),
    ("Framework", "Next.js", "Version 16 (React 19, App Router)"),
    ("Language", "TypeScript", "Version 5.x"),
    ("Database", "PostgreSQL", "Version 15 or later"),
    ("ORM", "Prisma", "Version 5.x (schema-first ORM)"),
    ("AI Framework", "LangChain", "@langchain/google-genai, @langchain/core"),
    ("AI Model", "Google Gemini", "gemini-2.0-flash or gemini-1.5-pro"),
    ("Authentication", "JWT + bcryptjs", "jsonwebtoken 9.x, bcryptjs 2.x"),
    ("UI Components", "Radix UI + Tailwind CSS", "Latest stable versions"),
    ("Chrome Extension", "Manifest V3", "Chrome 90+ (MV3 support)"),
    ("IDE / Editor", "VS Code", "With ESLint and Prettier extensions"),
    ("Package Manager", "npm", "Version 10.x"),
]
add_table(["Component", "Technology / Tool", "Version / Notes"], sw_data, "2.2", "Software Requirements")

section_heading("2.4", "USE CASE DIAGRAM")
body("The Use Case Diagram illustrates the functional interactions between the two primary actors — the Authenticated User and the System (PhishGuardAI) — and all the use cases supported by the phishing detection application. The diagram captures the complete scope of functionality accessible through both the Chrome extension and the web-based admin dashboard.")
add_para("")
body("Actors: (1) User (Authenticated): An individual who has registered and logged in through the Chrome extension. The user's browsing is automatically monitored and protected. (2) Administrator: A user with the ADMIN role who has access to the web-based dashboard for monitoring and analysis. (3) External Services: Google Gemini AI (for phishing analysis) and PostgreSQL Database.")
add_para("")
body("Primary Use Cases: Register Account, Login/Authenticate, Auto-Scan URL on Navigation, View Phishing Block Page, View Warning Page, Proceed Despite Warning, Check URL Manually via Popup, Configure Extension Settings, View Admin Dashboard, Filter/Search URL History, View Audit Records, Export/Analyze Data.")

img_placeholder("2.1", "Use Case Diagram — PhishGuardAI System")

section_heading("2.5", "DATA FLOW DIAGRAM")
subsection_heading("2.5.1", "Level 0 – Context Diagram")
body("At the context level (Level 0), PhishGuardAI is represented as a single system that interacts with three external entities: the User (providing URLs through browsing activity and authentication credentials, receiving protection actions), Google Gemini AI (receiving URL and content analysis requests, returning classification results with confidence scores), and the PostgreSQL Database (receiving analysis results for caching, returning cached results for known URLs).")

img_placeholder("2.2", "Data Flow Diagram – Level 0 (Context Diagram)")

subsection_heading("2.5.2", "Level 1 – Detailed Data Flow")
body("At Level 1, the system is decomposed into five major sub-processes: (1) User Authentication and Management — handles registration, login, JWT token issuance and validation, and role-based access control; (2) URL Interception and Normalization — extracts base domain from navigated URLs and normalizes for consistent caching; (3) Cache Lookup and Management — checks in-memory and database caches, stores new results; (4) AI-Powered Analysis — orchestrates LangChain tool calling with Gemini for URL and content analysis; (5) Action Execution and Audit Logging — applies blocking, warning, or allow action and records all security events with user context.")

img_placeholder("2.3", "Data Flow Diagram – Level 1 (Detailed Internal Flow)")

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — SYSTEM DESIGN
# ════════════════════════════════════════════════════════════════════════════
chapter_heading("3", "SYSTEM DESIGN")

section_heading("3.1", "PROPOSED SYSTEM")
body("The proposed PhishGuardAI system represents a significant advancement over existing phishing protection solutions by combining real-time AI reasoning with efficient distributed caching and a comprehensive audit framework. The system's proposed design addresses the limitations of existing tools through the following architectural and functional innovations:")
bullet("Semantic AI Analysis: Unlike rule-based and ML classifiers that rely on pre-computed features, the LangChain-Gemini integration reasons about URL and website legitimacy using natural language understanding — the same cognitive process a trained security analyst would apply when manually investigating a suspicious site.")
bullet("Tool-Augmented Analysis: When URL pattern analysis alone yields insufficient confidence (below 0.8 threshold), the AI automatically invokes the fetchWebsiteContent tool to inspect the actual webpage — including HTML structure, JavaScript code, form submission targets, and link patterns — providing evidence-based analysis rather than inference from URL features alone.")
bullet("Shared Threat Intelligence: Every URL analyzed for any user is stored in a shared cache, creating a collective intelligence network where threats discovered by one user immediately protect all other users without additional API calls or processing costs.")
bullet("Nuanced Risk Categorization: The five-category warning system (PIRACY, SCAMMING, RISKY_LINKS, SCAM_PRODUCTS, RISKY_FILES) enables differentiated responses that inform users about the specific nature of a risk rather than applying a generic 'dangerous' label, improving user trust and decision-making.")
bullet("Open and Accessible: The system operates as a browser extension with a self-hostable backend, requiring only a Google Gemini API key and a PostgreSQL database — both of which have free tiers — making professional-quality AI security accessible without per-query subscription costs.")
bullet("Fail-Open Security Model: The system prioritizes user experience by failing open (allowing navigation) on errors rather than blocking legitimate browsing when the service is unavailable, while maintaining complete audit records of all security events.")

section_heading("3.2", "SYSTEM ARCHITECTURE")
subsection_heading("3.2.1", "Architecture Overview")
body("The complete system architecture consists of four distinct tiers:")
bullet("Presentation Tier (Chrome Extension): The Manifest V3 Chrome extension serves as the user-facing interface. It consists of a background service worker for URL interception and cache management, content scripts for warning banner injection, popup.html/js for manual URL checking, options.html/js for settings management, scanning.html for the analysis progress page, blocked.html for confirmed phishing sites, and warning.html for risky site warnings.")
bullet("Application Tier (Next.js 16 API): The backend API server provides RESTful endpoints for all extension operations. Built with Next.js 16's App Router and API Routes, it handles authentication with JWT and bcrypt, URL analysis with LangChain-Gemini integration, database operations through Prisma ORM, and admin dashboard functionality.")
bullet("Data Tier (PostgreSQL + Prisma): The PostgreSQL database serves as the persistent data store for URL analysis results, user accounts, and audit records. Prisma 5 provides type-safe query generation and schema management through migration files, ensuring consistent database state across deployments.")
bullet("External AI Tier (Google Gemini): Google's Gemini model, accessed through the LangChain framework, provides the semantic intelligence that powers phishing detection. The LangChain integration enables multi-turn analysis with structured tool calling for website content inspection.")

subsection_heading("3.2.2", "Architecture Layers Detail")
arch_data = [
    ("Presentation Layer", "Chrome Extension (MV3)", "Background service worker, content scripts, popup, options, scanning, blocked, warning pages"),
    ("API Gateway", "Next.js 16 API Routes", "URL checking, auth, admin endpoints, Prisma ORM integration"),
    ("AI Engine", "LangChain + Google Gemini", "Multi-turn analysis, tool calling, confidence scoring, warning categorization"),
    ("Data Layer", "PostgreSQL + Prisma 5", "URL cache, user accounts, audit records, indexed queries"),
    ("Auth Layer", "JWT + bcryptjs", "Stateless token auth, password hashing, role-based access"),
    ("Content Analysis", "fetchWebsiteContent Tool", "HTML parsing, JS analysis, form inspection, link extraction"),
]
add_table(["Layer", "Technology", "Details"], arch_data, "3.1A", "System Architecture Layer Details")

img_placeholder("3.1", "System Architecture Diagram")

section_heading("3.3", "REPOSITORY STRUCTURE")
body("The project is organized as a single Next.js repository with a clear separation of concerns between the backend API and the Chrome extension:")
subsection_heading("3.3.1", "Next.js Application (/src)")
body("The /src/app directory contains all API routes and pages using the App Router architecture. The /src/app/api directory houses the backend endpoints: /api/auth/login and /api/auth/register for authentication, /api/check-url for the main phishing detection endpoint, and /api/admin/urls and /api/admin/audit for admin dashboard data. The /src/app/admin directory contains the React admin dashboard pages and components.")
add_para("")
body("The /src/lib directory contains all shared utility code: gemini.ts and langchain-phishing.ts for AI integration, auth.ts for JWT and bcrypt utilities, middleware.ts for authentication middleware, prisma.ts for the database client singleton, and /src/lib/tools/fetch-website-content.ts for the LangChain website content analysis tool.")

subsection_heading("3.3.2", "Chrome Extension (/chrome-extension)")
body("The chrome-extension directory contains all extension source files. The manifest.json defines the extension's permissions, service worker, content scripts, and web-accessible resources. background.js implements the service worker with URL interception and cache management. content-scan.js injects warning banners into web pages. The popup, options, scanning, blocked, and warning HTML/JS pairs provide the user interface for each interaction type. api-client.js handles all API communication with the backend.")

subsection_heading("3.3.3", "Database (/prisma)")
body("The prisma directory contains the schema.prisma file defining all database models and relationships, the migrations directory with versioned SQL migration files, and the generated Prisma client in the generated subdirectory.")

img_placeholder("3.2", "Project Repository Structure Diagram")

section_heading("3.4", "API DESIGN")
body("The API follows RESTful conventions with consistent response formats and proper HTTP status codes. All protected routes require a Bearer JWT token in the Authorization header.")
api_data = [
    ("Auth", "POST", "/api/auth/login", "Public", "Authenticate user, return JWT token"),
    ("Auth", "POST", "/api/auth/register", "Public", "Create new user account"),
    ("URL Check", "POST", "/api/check-url", "USER", "Analyze URL for phishing/warnings"),
    ("Admin URLs", "GET", "/api/admin/urls", "ADMIN", "Paginated URL analysis history"),
    ("Admin Audit", "GET", "/api/admin/audit", "ADMIN", "Paginated user audit records"),
]
add_table(["Module", "Method", "Endpoint", "Access", "Purpose"], api_data, "3.1", "API Endpoint Summary")

body("The /api/check-url endpoint accepts a JSON body with a 'url' field. It returns: { isPhishing: boolean, hasWarning: boolean, warningType: string[], warningSeverity: string, confidence: number, reason: string, cached: boolean }. For cached results, the response is immediate; for new URLs, the endpoint invokes the LangChain analysis pipeline which may take 3-15 seconds depending on whether website content fetching is required.")

section_heading("3.5", "DATABASE SCHEMA")
subsection_heading("3.5.1", "Entity Relationships")
er_data = [
    ("User → UrlCheck", "One-to-Many (1:N)", "One user can have multiple URL checks associated via optional userId"),
    ("User → AuditRecord", "One-to-Many (1:N)", "One user can have multiple audit records of their browsing actions"),
    ("UrlCheck (standalone)", "Independent", "URL checks are cached independently and may exist without a user reference"),
]
add_table(["Relationship", "Type", "Description"], er_data, "3.2", "Database Entity Relationships")

img_placeholder("3.3", "Entity Relationship (ER) Diagram")

subsection_heading("3.5.2", "Data Models")
body("User Model: Stores authentication and profile information. Fields: id (UUID), email (unique), passwordHash (bcrypt), role (USER|ADMIN), createdAt, updatedAt. Relations: urlChecks[], auditRecords[].")
add_para("")
body("UrlCheck Model: Stores URL analysis results as a shared cache. Fields: id (UUID), url (unique indexed), isPhishing (boolean), hasWarning (boolean), confidence (float), warningType (string[]), warningSeverity (string), reason (AI explanation text), checkedAt, createdAt, updatedAt, userId (optional FK). The url field has a unique index enabling O(log n) lookups for cache hits.")
add_para("")
body("AuditRecord Model: Stores all user interactions with blocked or warned sites. Fields: id (UUID), userId (FK, required), url, isPhishing, hasWarning, action (BLOCKED|ALLOWED|WARNING_SHOWN), ipAddress, userAgent, visitedAt, createdAt. Indexed on userId, visitedAt, and action for efficient admin dashboard queries.")

img_placeholder("3.4", "Prisma Database Schema Screenshot")

section_heading("3.6", "LIST OF MODULES")
body("PhishGuardAI is organized into seven functional modules, each handling a specific aspect of the overall protection pipeline:")
modules_data = [
    ("User Authentication Module", "Handles user registration, login, JWT token management, and role-based authorization. Implements bcrypt password hashing and Prisma user operations.", "Next.js API Routes"),
    ("URL Interception Module", "Chrome extension service worker that monitors tab navigation events, normalizes URLs, manages in-memory cache, and routes requests through the protection pipeline.", "Chrome Extension"),
    ("AI Analysis Engine", "Orchestrates LangChain-Gemini integration for phishing detection. Manages multi-turn analysis, tool calling, JSON response parsing, and confidence scoring.", "Next.js + LangChain"),
    ("Website Content Analyzer", "LangChain tool that fetches and parses website HTML, extracts scripts, forms, links, and suspicious patterns. Provides structured content data to the AI model.", "Next.js + Cheerio"),
    ("Cache Management Module", "Manages the shared PostgreSQL URL cache and extension in-memory cache. Handles URL normalization, cache invalidation, and checkOnly queries for fast lookups.", "Next.js + Prisma"),
    ("Admin Dashboard Module", "Provides React-based admin interface with paginated URL history, audit records, search and filtering. JWT authentication required with ADMIN role.", "Next.js React"),
    ("Audit Logging Module", "Records all user interactions with phishing and warning sites. Captures user identity, URL, action taken, IP address, user agent, and timestamp.", "Next.js + Prisma"),
]
add_table(["Module Name", "Description", "Technology"], modules_data, "3.3", "List of Modules")

section_heading("3.7", "OVERALL SYSTEM DESIGN")
subsection_heading("3.7.1", "Input Design")
body("The system accepts two types of input. Automatic input occurs when the Chrome extension intercepts URL navigation through the chrome.tabs.onUpdated event — the user takes no explicit action. The URL is automatically extracted, normalized (removing paths, query parameters, and fragments to yield the base domain), and processed through the protection pipeline. Manual input occurs when the user opens the extension popup and enters a URL directly for on-demand analysis, or when an administrator uses the dashboard search to look up specific URLs or users.")

subsection_heading("3.7.2", "Processing Design")
body("The processing pipeline follows a short-circuit evaluation pattern optimized for performance. For each URL, the system first checks the in-memory extension cache (O(1) hash lookup, sub-millisecond). If not in memory, it queries the database with checkOnly=true, which performs a SELECT on the indexed url column (typically under 10ms). Only when the URL is not found in either cache does the full AI analysis begin.")
add_para("")
body("The AI analysis invokes the LangChain agent with the URL and a detailed system prompt establishing the role of a cybersecurity expert analyst. The Gemini model performs initial analysis from the URL alone; if confidence is below the 0.8 threshold, it invokes the fetchWebsiteContent tool. This tool makes an HTTP GET request to the URL, parses the HTML with regex-based extraction (to avoid executing JavaScript), extracts forms, scripts, links, and suspicious patterns, and returns a structured JSON summary to the AI model. The model then produces its final classification.")

subsection_heading("3.7.3", "Output Design")
body("The system produces three types of output. For confirmed phishing sites (isPhishing=true), the extension redirects the tab to blocked.html, passing the phishing URL and AI reasoning as query parameters. The blocked page displays the URL, the threat explanation, and options to go back or report a false positive. For risky sites (hasWarning=true), the extension either redirects to warning.html (if the 'block warnings' setting is enabled) or injects a warning banner into the page via the content script. For safe sites, navigation proceeds transparently with no user-visible action.")

section_heading("3.8", "SYSTEM TESTING")
body("System testing for PhishGuardAI was conducted across multiple dimensions to validate all functional requirements, integration points, and edge cases.")

subsection_heading("3.8.1", "Unit Testing")
body("Unit tests were written for individual service functions. The authentication service was tested with valid credentials, invalid passwords, duplicate email registration, and malformed JWT tokens. The URL normalization function was tested with various URL formats including URLs with and without paths, query strings, fragments, 'www' prefixes, and different protocols. The LangChain analysis module was tested with known phishing URLs (from PhishTank), known safe URLs, and edge cases including localhost addresses and IP-based URLs.")

subsection_heading("3.8.2", "Integration Testing")
body("Integration testing verified the complete data flow through all system components. Key integration points tested include: the Chrome extension to Next.js API communication with valid and invalid JWT tokens; the Next.js API to Prisma database operations including cache writes and reads; the LangChain framework to Google Gemini API integration including tool calling and response parsing; the complete URL check pipeline from extension interception to database cache storage; and the admin dashboard API to PostgreSQL queries with pagination and filtering.")

subsection_heading("3.8.3", "Functional Testing")
body("Functional testing validated that all specified use cases were correctly implemented. Tests covered the complete user journey from extension installation through authentication, browsing protection, phishing blocking, warning display, and admin monitoring. All five warning categories were verified with appropriate test URLs. The fail-open behavior was verified by simulating API unavailability.")

subsection_heading("3.8.4", "Performance Testing")
body("Performance testing measured response times across different cache states. Cached URL responses (in-memory) were measured at under 5ms end-to-end. Database cache hits returned results in under 200ms including network round-trip. Full AI analysis for URL-only checks averaged 2-4 seconds. AI analysis with website content fetching averaged 5-12 seconds. Database query times for admin dashboard pagination averaged under 50ms for tables with 10,000+ records.")

section_heading("3.9", "TEST CASES")
subsection_heading("3.9.1", "Authentication Module Test Cases")
auth_tc = [
    ("TC-A01", "User registration with valid data", "POST /api/auth/register with name, email, password", "201 Created; user in DB; password hashed", "Pass"),
    ("TC-A02", "Registration with duplicate email", "POST /api/auth/register with existing email", "409 Conflict — Email already registered", "Pass"),
    ("TC-A03", "Login with valid credentials", "POST /api/auth/login with correct credentials", "200 OK; JWT token returned in response", "Pass"),
    ("TC-A04", "Login with wrong password", "POST /api/auth/login with incorrect password", "401 Unauthorized — Invalid credentials", "Pass"),
    ("TC-A05", "Access protected route without token", "GET /api/admin/urls (no Authorization header)", "401 Unauthorized — Token missing", "Pass"),
    ("TC-A06", "Access admin route as USER role", "GET /api/admin/urls with USER JWT", "403 Forbidden — Admin access required", "Pass"),
    ("TC-A07", "Expired JWT token", "Request with expired JWT", "401 Unauthorized — Token expired", "Pass"),
    ("TC-A08", "Valid admin login and dashboard access", "POST /api/auth/login as ADMIN, then GET /api/admin/urls", "200 OK with paginated URL data", "Pass"),
]
add_table(["TC ID", "Description", "Input / Action", "Expected Output", "Status"], auth_tc, "3.4", "Test Cases – Authentication Module")

subsection_heading("3.9.2", "URL Checking Module Test Cases")
url_tc = [
    ("TC-U01", "Check known phishing URL", "POST /api/check-url with known phishing domain", "isPhishing: true; confidence > 0.8; reason provided", "Pass"),
    ("TC-U02", "Check known safe URL", "POST /api/check-url with trusted domain (e.g. google.com)", "isPhishing: false; hasWarning: false", "Pass"),
    ("TC-U03", "Check previously cached URL", "POST /api/check-url with already-analyzed URL", "Same result returned; cached: true in response", "Pass"),
    ("TC-U04", "Check URL with warning (piracy site)", "POST /api/check-url with known piracy domain", "hasWarning: true; warningType includes PIRACY", "Pass"),
    ("TC-U05", "Check URL without authentication", "POST /api/check-url without JWT token", "401 Unauthorized", "Pass"),
    ("TC-U06", "Check malformed URL", "POST /api/check-url with 'not-a-url'", "400 Bad Request or AI returns low confidence", "Pass"),
    ("TC-U07", "URL normalization consistency", "Check 'https://example.com/path?q=1' and 'http://www.example.com'", "Both resolve to same cached result for example.com", "Pass"),
    ("TC-U08", "Scamming site detection", "POST /api/check-url with known scam site", "hasWarning: true; warningType includes SCAMMING", "Pass"),
]
add_table(["TC ID", "Description", "Input / Action", "Expected Output", "Status"], url_tc, "3.5", "Test Cases – URL Checking Module")

subsection_heading("3.9.3", "AI Analysis Module Test Cases")
ai_tc = [
    ("TC-AI01", "LangChain tool calling triggered", "Submit URL with low-confidence URL analysis", "fetchWebsiteContent tool invoked; deeper analysis performed", "Pass"),
    ("TC-AI02", "Confidence scoring accuracy", "Submit high-confidence phishing URL", "Confidence score > 0.9 returned", "Pass"),
    ("TC-AI03", "JSON response parsing", "AI returns valid JSON classification", "Structured result extracted; stored in DB", "Pass"),
    ("TC-AI04", "API error handling", "Simulate Gemini API timeout", "Fail-open: navigation allowed; error logged", "Pass"),
    ("TC-AI05", "Warning category assignment", "Submit risky-links site", "warningType: ['RISKY_LINKS']; severity assigned", "Pass"),
]
add_table(["TC ID", "Description", "Input / Action", "Expected Output", "Status"], ai_tc, "3.6", "Test Cases – AI Analysis Module")

subsection_heading("3.9.4", "Chrome Extension Module Test Cases")
ext_tc = [
    ("TC-E01", "URL interception on navigation", "Navigate to any URL in Chrome with extension active", "Extension intercepts; cache/API check initiated", "Pass"),
    ("TC-E02", "Phishing site blocked", "Navigate to confirmed phishing URL", "Tab redirected to blocked.html with URL and reason", "Pass"),
    ("TC-E03", "Warning page displayed", "Navigate to warning-category URL (block warnings enabled)", "Tab redirected to warning.html with category info", "Pass"),
    ("TC-E04", "Safe site navigation transparent", "Navigate to safe URL", "No extension interference; page loads normally", "Pass"),
    ("TC-E05", "Extension popup manual check", "Open popup; enter URL; click Check", "Result displayed in popup with risk assessment", "Pass"),
    ("TC-E06", "Auto-scan toggle functionality", "Disable auto-scan in options; navigate to phishing URL", "Extension does not intercept; auto-scan disabled", "Pass"),
    ("TC-E07", "In-memory cache hit", "Navigate to same URL twice in same session", "Second navigation instant (in-memory cache)", "Pass"),
    ("TC-E08", "Proceed from warning page", "Click 'Proceed Anyway' on warning.html", "Original URL opened; audit record created", "Pass"),
]
add_table(["TC ID", "Description", "Input / Action", "Expected Output", "Status"], ext_tc, "3.7", "Test Cases – Chrome Extension Module")

subsection_heading("3.9.5", "Admin Dashboard Module Test Cases")
admin_tc = [
    ("TC-AD01", "Admin login and redirect", "Login with ADMIN credentials in web app", "Redirect to /admin dashboard", "Pass"),
    ("TC-AD02", "Non-admin redirect to login", "Access /admin without ADMIN role token", "Redirect to /admin/login", "Pass"),
    ("TC-AD03", "URL table pagination", "GET /api/admin/urls?page=2&limit=10", "10 URL records from page 2 returned", "Pass"),
    ("TC-AD04", "Filter by phishing status", "GET /api/admin/urls?isPhishing=true", "Only phishing URLs returned", "Pass"),
    ("TC-AD05", "Audit records date filter", "GET /api/admin/audit?startDate=2024-01-01&endDate=2024-12-31", "Records within date range returned", "Pass"),
]
add_table(["TC ID", "Description", "Input / Action", "Expected Output", "Status"], admin_tc, "3.8", "Test Cases – Admin Dashboard Module")

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — IMPLEMENTATION AND ANALYSIS OF RESULTS
# ════════════════════════════════════════════════════════════════════════════
chapter_heading("4", "IMPLEMENTATION AND ANALYSIS OF RESULTS")

section_heading("4.1", "IMPLEMENTATION ENVIRONMENT")
body("PhishGuardAI is implemented using a modern full-stack TypeScript/JavaScript technology combination. The development environment was configured with Node.js 20 LTS as the runtime for the Next.js backend, PostgreSQL 15 as the relational database, and Google Chrome as the target browser for extension testing. The development workstation ran macOS with VS Code as the primary IDE, supplemented by the Chrome Extension Developer Tools for extension debugging and the Prisma Studio GUI for database inspection.")
add_para("")
body("The implementation follows a modular development approach where each functional module was developed and tested independently before integration. The Next.js backend uses TypeScript throughout with strict mode enabled, providing compile-time type safety for all API operations, database queries, and AI response handling. The Chrome extension is implemented in vanilla JavaScript to minimize extension size and avoid bundling complexity.")
add_para("")
body("Version control is managed through Git with a single repository structure. The project uses a GitFlow-inspired branch strategy with main (production-ready), develop (integration), and feature branches. Database schema changes are managed through Prisma migration files that provide a versioned history of all schema changes with rollback capability.")
add_para("")
body("Environment configuration uses a .env file for the development environment with the following required variables: DATABASE_URL (PostgreSQL connection string), GEMINI_API_KEY (Google AI Studio API key), GEMINI_MODEL_NAME (model identifier), and JWT_SECRET (minimum 32-character random string for token signing). A seed script (scripts/seed-admin.ts) creates the initial admin user with secure credentials.")

section_heading("4.2", "TECHNOLOGY STACK")
body("The complete technology stack for PhishGuardAI is summarized in the following table:")
tech_data = [
    ("Backend Framework", "Next.js 16", "API Routes with App Router; TypeScript; server-side rendering for admin dashboard"),
    ("AI Framework", "LangChain (@langchain/google-genai)", "Multi-turn agent execution, tool calling, structured output parsing"),
    ("AI Model", "Google Gemini (gemini-2.0-flash)", "Large language model for semantic phishing analysis"),
    ("Database", "PostgreSQL 15", "Relational database for URL cache, user accounts, audit records"),
    ("ORM", "Prisma 5", "Type-safe database queries, schema management, migrations"),
    ("Authentication", "jsonwebtoken + bcryptjs", "Stateless JWT auth (7-day expiry), bcrypt password hashing"),
    ("Content Parsing", "Cheerio / Regex", "HTML parsing for website content analysis tool"),
    ("UI Components", "Radix UI + Tailwind CSS", "Accessible component primitives for admin dashboard"),
    ("Chrome Extension", "Manifest V3 (Vanilla JS)", "Service worker, content scripts, declarativeNetRequest"),
    ("Schema Validation", "Zod", "Runtime validation for LangChain tool definitions and API inputs"),
    ("Runtime", "Node.js 20 LTS", "JavaScript runtime for Next.js server"),
    ("Package Manager", "npm 10", "Dependency management for Next.js application"),
]
add_table(["Component", "Technology", "Purpose"], tech_data, "4.1", "Technology Stack")

section_heading("4.3", "MODULES COMPLETED")
body("The following sections describe the implementation details of each module completed as part of the PhishGuardAI system.")

subsection_heading("4.3.1", "User Authentication Module")
body("The authentication module is implemented across two Next.js API route files: /src/app/api/auth/login/route.ts and /src/app/api/auth/register/route.ts, with shared utilities in /src/lib/auth.ts.")
add_para("")
body("The registration endpoint validates the email format using a regex pattern and enforces a minimum password length of 8 characters. It checks for duplicate email addresses using a Prisma unique constraint query before hashing the password with bcryptjs using 10 salt rounds. The resulting user record is created in PostgreSQL with the USER role by default. Administrators can be created using the seed script, which sets the ADMIN role.")
add_para("")
body("The login endpoint retrieves the user record by email, then uses bcryptjs.compare() to verify the submitted password against the stored hash. On success, it generates a JWT token signed with the JWT_SECRET environment variable and a 7-day expiry. The token payload includes the user's id and email. The /src/lib/middleware.ts module provides requireAuth() and requireAdmin() higher-order functions that validate the Bearer token and optionally check the ADMIN role for all protected API routes.")

img_placeholder("4.1", "User Authentication Flow (Login and Register Interface)")

subsection_heading("4.3.2", "URL Checking and Caching Module")
body("The URL checking module is implemented in /src/app/api/check-url/route.ts. This is the most critical endpoint in the system, handling all phishing analysis requests from the Chrome extension.")
add_para("")
body("Upon receiving a POST request with a URL, the endpoint first verifies the JWT token using requireAuth(). It then normalizes the URL by extracting the base domain using the JavaScript URL constructor and removing common tracking parameters. The normalized URL is used as the cache key for all database operations.")
add_para("")
body("The caching strategy proceeds in order: First, the endpoint queries the UrlCheck table for an existing record with the normalized URL. If found (cache hit), the stored result is returned immediately with cached:true in the response. For cache misses, the LangChain analysis pipeline is invoked. After analysis, the result is stored in the UrlCheck table using Prisma's upsert operation (creating or updating based on URL uniqueness). An AuditRecord is also created if the result is phishing or has a warning, capturing the user's context.")

img_placeholder("4.2", "Chrome Extension Popup Interface")

subsection_heading("4.3.3", "LangChain AI Analysis Engine")
body("The AI analysis engine is implemented in /src/lib/langchain-phishing.ts. This module orchestrates the multi-turn conversation with Google Gemini using LangChain's ChatGoogleGenerativeAI model.")
add_para("")
body("The analysis begins by constructing a HumanMessage containing the URL to analyze along with a detailed system prompt that establishes the AI's role as an expert cybersecurity analyst. The prompt instructs the model to consider: domain name patterns (typosquatting, homograph attacks), URL structural anomalies, SSL/TLS certificate patterns, known phishing indicators, and website content analysis (if needed).")
add_para("")
body("The model is initialized with temperature=0.1 for consistent, deterministic results and bound to the fetchWebsiteContent tool using LangChain's bindTools() method. The analysis loop runs for up to 5 iterations, allowing the model to call the tool, observe the results, and continue reasoning. The final response is parsed from JSON using a regex extraction fallback to handle cases where the model wraps the JSON in markdown code blocks.")
add_para("")
body("The structured response follows this schema: { isPhishing: boolean, hasWarning: boolean, warningType: string[], warningSeverity: 'low'|'medium'|'high', confidence: number, reason: string }.")

img_placeholder("4.8", "LangChain AI Analysis Flow Diagram")

subsection_heading("4.3.4", "Website Content Analysis Tool")
body("The fetchWebsiteContent LangChain tool is implemented in /src/lib/tools/fetch-website-content.ts. This tool extends LangChain's StructuredTool base class and is defined with a Zod schema requiring a 'url' parameter.")
add_para("")
body("When invoked by the AI, the tool makes an HTTP GET request to the URL using Node.js fetch, with a 10-second timeout and a realistic User-Agent header to avoid bot detection. The HTML response is processed through regex-based extraction (rather than a full DOM parser) to identify: the page title from the <title> tag, meta description from the appropriate meta element, all inline <script> content and external script src attributes, form action URLs and input field names, all anchor href attributes classified as internal vs external links, and suspicious patterns including base64 encoding, eval() calls, and document.write() usage.")
add_para("")
body("The extracted data is assembled into a structured object and serialized to JSON. To stay within the Gemini model's context window limits, the tool applies token-aware trimming that truncates JavaScript content at 500 characters per script while preserving all structural elements. The final output is returned as a formatted string to the AI model, which uses this evidence to finalize its classification.")

subsection_heading("4.3.5", "Chrome Extension Implementation")
body("The Chrome extension is implemented as a Manifest V3 extension with a service worker background script, content scripts, and multiple HTML interface pages.")
add_para("")
body("The background service worker (background.js) registers a chrome.tabs.onUpdated listener that fires whenever a tab navigates to a new URL. The handler checks if the URL uses the http or https protocol (ignoring chrome:// and extension pages), extracts the base domain for cache lookup, and checks the in-memory JavaScript Map cache. If not cached, it calls the backend API with checkOnly=true first — a fast database-only lookup that avoids triggering AI analysis — and if not found in the database cache, it redirects the tab to the scanning.html page, which then performs the full analysis.")
add_para("")
body("The scanning page (scanning.js) shows a loading animation while calling /api/check-url with the original URL and the stored JWT token. Based on the result, it either redirects to blocked.html (phishing), warning.html (warning with block-warnings setting), injects a content script banner (warning without block-warnings), or navigates directly to the original URL (safe).")

img_placeholder("4.3", "URL Scanning Page — Loading Animation During Analysis")
img_placeholder("4.4", "Phishing Blocked Page — Threat Details and Back Button")
img_placeholder("4.5", "Warning Page — Risk Category and Proceed/Back Options")

subsection_heading("4.3.6", "Admin Dashboard Module")
body("The admin dashboard is implemented as a Next.js React application in /src/app/admin/. The main page component uses React's useState and useEffect hooks to manage authentication state and tab selection. On mount, it validates the stored JWT token from localStorage; if invalid or absent, it redirects to /admin/login.")
add_para("")
body("The URL table component (/src/app/admin/components/url-table.tsx) fetches paginated data from /api/admin/urls with support for query parameters: page, limit, search (URL search), isPhishing filter, and hasWarning filter. The table displays each URL's analysis result, confidence score as a percentage, warning category and severity, and the timestamp of analysis. The audit table component provides similar functionality for /api/admin/audit with additional filters for date range and action type (BLOCKED, ALLOWED, WARNING_SHOWN).")

img_placeholder("4.6", "Admin Dashboard — URL Checks Table with Filters")
img_placeholder("4.7", "Admin Dashboard — Audit Records Table")

subsection_heading("4.3.7", "Audit Logging Module")
body("Audit logging is integrated directly into the URL check pipeline. Whenever a URL analysis result is phishing (isPhishing=true) or triggers a warning (hasWarning=true), an AuditRecord is created in the PostgreSQL database through Prisma. The record captures the authenticated user's ID (from the validated JWT), the analyzed URL, the classification result, the action that will be taken by the extension (BLOCKED, WARNING_SHOWN), the IP address extracted from the X-Forwarded-For or remote address headers, the User-Agent string, and the current timestamp.")
add_para("")
body("For safe URLs, no audit record is created by default, as recording every safe page visit would generate excessive data and potential privacy concerns. Only security-relevant events (blocks and warnings) are recorded, providing a focused audit trail for security analysis without unnecessary data collection.")

section_heading("4.4", "DEVIATIONS AND JUSTIFICATIONS")
body("During implementation, several deviations from the initially conceived approach were made. These deviations were driven by technical constraints, API limitations, and opportunities to improve system quality.")
dev_data = [
    ("D-01", "Direct Gemini API integration (without LangChain)", "LangChain ChatGoogleGenerativeAI with tool calling", "LangChain's multi-turn agent pattern and structured tool calling enables the AI to request website content analysis only when needed, reducing API costs for simple phishing URLs while enabling deep analysis for complex cases."),
    ("D-02", "Single-step URL analysis only", "Two-step analysis: URL pattern first, website content fetch only when confidence < 0.8", "The two-step approach reduces latency for high-confidence phishing URLs (typosquatting, known patterns) that don't require content inspection, while maintaining thoroughness for ambiguous cases."),
    ("D-03", "Separate Chrome extension and backend repositories", "Monorepo: chrome-extension directory inside Next.js project", "A single repository simplifies development, ensures API contract consistency, and allows shared environment configuration. Extension and API are deployed independently but developed together."),
    ("D-04", "Session storage for JWT in extension", "chrome.storage.local for persistent JWT storage", "chrome.storage.local persists authentication across browser sessions and extension restarts, providing a better user experience than session-only storage that requires re-login after every browser restart."),
    ("D-05", "Separate Python FastAPI for NLP processing", "All processing in Next.js API routes with LangChain", "The phishing detection task maps naturally to LLM analysis rather than traditional NLP (LSA/YAKE). Using LangChain within Next.js eliminates a separate service, simplifying deployment while achieving superior detection quality."),
    ("D-06", "Real-time WebSocket for analysis status", "Redirect to scanning.html with polling fallback", "WebSocket implementation adds significant complexity for a use case that resolves in under 15 seconds. The redirect pattern provides a clean, progress-visible loading experience without WebSocket infrastructure."),
]
add_table(["#", "Planned Design", "Actual Implementation", "Justification"], dev_data, "4.3", "Deviations and Justifications")

section_heading("4.5", "PROJECT ROADMAP AND STATUS")
body("The project development was organized into six phases with clearly defined deliverables. The following table summarizes the planned phases, their activities, and current status:")
road_data = [
    ("Phase 1", "Requirements Analysis and System Design: Defined functional and non-functional requirements, designed system architecture, database schema, and API contracts. Established technology stack decisions.", "Completed"),
    ("Phase 2", "Core Backend Development and Authentication: Implemented Next.js API server, Prisma schema, PostgreSQL setup, JWT authentication module, and user management endpoints.", "Completed"),
    ("Phase 3", "AI Integration and URL Analysis Pipeline: Integrated LangChain with Google Gemini, implemented fetchWebsiteContent tool, developed confidence scoring system, implemented database caching.", "Completed"),
    ("Phase 4", "Chrome Extension Development: Implemented Manifest V3 service worker, URL interception, in-memory caching, scanning/blocked/warning pages, popup and options interfaces.", "Completed"),
    ("Phase 5", "Admin Dashboard Development: Built React admin dashboard with URL history table, audit records table, pagination, search and filtering, and admin authentication.", "Completed"),
    ("Phase 6", "Testing, Refinement, and Documentation: Conducted unit, integration, functional, and performance testing. Addressed identified issues. Prepared project documentation.", "Completed"),
]
add_table(["Phase", "Activity", "Status"], road_data, "4.4", "Project Roadmap and Phase Status")

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 — CONCLUSION AND FUTURE WORK
# ════════════════════════════════════════════════════════════════════════════
chapter_heading("5", "CONCLUSION AND FUTURE WORK")

section_heading("5.1", "LIMITATIONS")
body("While PhishGuardAI successfully achieves its primary objectives, several limitations should be acknowledged for a complete understanding of the system's current capabilities and boundaries:")
bullet("Dependency on External AI Services: The system's detection quality depends entirely on the Google Gemini API. Service outages, API changes, or pricing model changes could affect availability and cost. The API has usage quotas and rate limits that may constrain high-volume deployments. An internet connection is required for all new URL analyses, with no offline fallback.")
bullet("Analysis Latency for Novel URLs: For URLs not in the cache, full AI analysis can take 5-15 seconds when website content fetching is required. While this is acceptable as a one-time cost per domain, users may find the scanning page delay frustrating for legitimate sites they trust. The analysis time is outside the system's control as it depends on Gemini API response time and the target website's loading speed.")
bullet("Chrome-Only Support: The current implementation supports only Google Chrome with Manifest V3. Users of Firefox, Safari, Edge, or Brave browsers do not receive the phishing protection offered by the extension. While the Next.js backend API is browser-agnostic, the extension itself would require porting to each browser's extension API.")
bullet("English-Language Optimization: The AI prompt and phishing pattern recognition are optimized for English-language websites. Phishing sites targeting non-English speakers in languages such as Tamil, Hindi, French, or Spanish may receive less accurate classifications, as the AI's pattern recognition for deceptive language is strongest in English.")
bullet("Cache Staleness Risk: The URL cache does not implement automatic expiry or freshness checking. A legitimate domain that is later compromised and turned into a phishing site may be cached as 'safe' from a previous analysis. Administrators must manually clear the UrlCheck record to force re-analysis of a cached domain.")
bullet("In-Memory Cache Volatility: The extension's in-memory Map cache is lost when the service worker is suspended by Chrome (typically after 5 minutes of inactivity). This means that URLs cached in memory during one browsing session may require database lookups on subsequent navigations rather than instant memory responses.")
bullet("No Email Phishing Detection: The system exclusively protects against web-based phishing encountered through browser navigation. Phishing links arriving via email, SMS, or messaging applications are not intercepted by the Chrome extension's URL interception mechanism.")

section_heading("5.2", "FUTURE ENHANCEMENT")
body("Building on the current implementation, several significant enhancements are proposed for future development:")
bullet("Multi-Browser Support: Porting the Chrome extension to Firefox (Manifest V2/V3 compatible), Microsoft Edge, and Safari would dramatically expand the user base. The Next.js backend API requires no changes for multi-browser support, as the extension communicates via standard HTTPS REST calls.")
bullet("Cache Expiry and Re-Analysis: Implementing a configurable TTL (Time-To-Live) on UrlCheck records — with shorter expiry (24 hours) for borderline cases and longer expiry (30 days) for high-confidence results — would address the cache staleness risk. Automatic re-analysis of expired records could be implemented as a background job using a cron scheduler.")
bullet("Real-Time WebSocket Status Updates: Replacing the scanning.html polling approach with a WebSocket connection using Socket.io would allow users to see live analysis progress — 'Analyzing URL structure...', 'Fetching website content...', 'Evaluating security indicators...' — improving the user experience for longer analysis tasks.")
bullet("Threat Intelligence Sharing API: Exposing an API endpoint for sharing anonymized threat intelligence with the broader security community (compatible with STIX/TAXII formats) would allow PhishGuardAI to both contribute to and benefit from collective phishing intelligence networks.")
bullet("Fine-Tuned Detection Model: Training a specialized phishing detection model on a curated dataset of confirmed phishing examples would enable faster, more cost-effective detection for common patterns, with the Gemini LLM reserved for novel or complex cases where general reasoning is required.")
bullet("OAuth 2.0 Authentication: Adding Google and Microsoft OAuth login options would reduce friction for users who prefer social login over email/password authentication. OAuth integration would also enable team-based deployments where organization members automatically receive admin access.")
bullet("Mobile Browser Extension: Developing a companion application for iOS and Android that integrates with Safari's content blocking API and Chrome's mobile extension support would extend protection to mobile browsing.")
bullet("API Rate Limiting and Quota Management: Implementing per-user API rate limiting, request queuing, and graceful degradation when approaching Gemini API quotas would improve system reliability in high-volume deployments.")
bullet("False Positive Reporting: Adding a user feedback mechanism allowing users to report incorrect classifications would create a labeled dataset for model improvement and enable administrators to review and override AI decisions.")
bullet("Integration with Security Awareness Training Platforms: Developing integrations with platforms like KnowBe4, Proofpoint Security Awareness, or Mimecast would allow enterprises to combine real-time phishing blocking with user education triggered when users encounter phishing attempts.")

section_heading("5.3", "CONCLUSION")
body("The PhishGuardAI project successfully demonstrates the feasibility and practical value of deploying Large Language Model intelligence as a real-time cybersecurity tool in a widely-used consumer application context. By integrating Google Gemini LLM through the LangChain framework into a Chrome browser extension with a Next.js backend, the project bridges the gap between advanced AI research capabilities and practical, accessible security protection.")
add_para("")
body("The system's core technical contribution is the multi-turn, tool-augmented analysis pipeline that enables the AI to reason about phishing indicators from both URL patterns and live website content — the first such approach deployed in a complete, end-to-end browser security system. Unlike supervised ML classifiers that require labeled training data and periodic retraining, the LLM-based approach generalizes to novel attack patterns through natural language reasoning, providing a fundamentally different and more adaptive detection mechanism.")
add_para("")
body("The three-level caching architecture (in-memory → database → AI analysis) demonstrates that AI-powered security can be made practical for real-time browser use despite the inherent latency of LLM inference. The shared database cache creates a network effect where the system's effectiveness grows with each analyzed URL, and the cost per analysis decreases as the cache hit rate increases over time.")
add_para("")
body("The comprehensive admin dashboard and audit logging infrastructure address the operational requirements of security-conscious organizations, providing the visibility and accountability needed for compliance and incident response. The role-based access control, JWT authentication, and bcrypt password hashing ensure that the system itself meets the security standards it is designed to enforce.")
add_para("")
body("From a technical evaluation perspective, the completed system correctly handles all tested security scenarios across authentication, URL analysis, phishing detection, warning categorization, and administrative monitoring. The fail-open error handling policy ensures that system reliability issues never inadvertently block legitimate user activity, prioritizing user experience without compromising the core security mission.")
add_para("")
body("The deviations from the original design specification — most notably the adoption of LangChain's agent framework over direct API calls, and the two-step analysis approach that reserves content fetching for low-confidence cases — represent informed engineering trade-offs that improve system quality, reduce operational costs, and enhance maintainability. These decisions demonstrate the importance of iterative refinement during implementation.")
add_para("")
body("In conclusion, PhishGuardAI makes a meaningful contribution to the intersection of AI-powered security tools and practical cybersecurity protection. The project demonstrates how thoughtful application of modern LLM capabilities, combined with sound software engineering principles and a user-centered design approach, can address a genuine and growing threat to digital security. The knowledge and skills acquired through this project provide a strong foundation for future work in AI-augmented security operations, browser extension development, and the broader field of cybersecurity tool design.")

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
page_break()
add_para("REFERENCES", WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, space_before=50, space_after=18)
refs = [
    '[1] Verizon. (2023). "2023 Data Breach Investigations Report." Verizon Business. Available at: https://www.verizon.com/business/resources/reports/dbir/',
    '[2] FBI Internet Crime Complaint Center (IC3). (2023). "2023 Internet Crime Report." Federal Bureau of Investigation.',
    '[3] Zhang, Y., Hong, J. I., and Cranor, L. F. (2007). "CANTINA: A Content-Based Approach to Detecting Phishing Web Sites." Proceedings of the 16th International World Wide Web Conference (WWW), 639-648.',
    '[4] Sahingoz, O. K., et al. (2019). "Machine Learning Based Phishing Detection from URLs." Expert Systems with Applications, 117, 345-357.',
    '[5] Mohammad, R. M., Thabtah, F., and McCluskey, L. (2015). "Tutorial and Critical Analysis of Phishing Websites Methods." Computer Science Review, 17, 1-24.',
    '[6] Brown, T. B., et al. (2020). "Language Models are Few-Shot Learners." Advances in Neural Information Processing Systems (NeurIPS), 33, 1877-1901.',
    '[7] Motlagh, F. N., et al. (2024). "Large Language Models in Cybersecurity: State-of-the-Art." arXiv preprint arXiv:2402.00891.',
    '[8] Chase, H. (2022). "LangChain: Building Applications with LLMs through Composability." GitHub. Available at: https://github.com/langchain-ai/langchain',
    '[9] Carlini, N., et al. (2012). "An Evaluation of the Google Chrome Extension Security Architecture." Proceedings of the 21st USENIX Security Symposium, 97-111.',
    '[10] Jagpal, N., et al. (2015). "Trends and Lessons from Three Years Fighting Malicious Extensions." Proceedings of the 24th USENIX Security Symposium, 579-593.',
    '[11] Weaver, N., et al. (2016). "Redirections, Befriending, and the Art of Email Phishing." Proceedings of the ACM Internet Measurement Conference.',
    '[12] Google. (2024). "Gemini API Documentation." Google AI for Developers. Available at: https://ai.google.dev/docs',
    '[13] LangChain. (2024). "LangChain Documentation — Google Generative AI Integration." Available at: https://js.langchain.com/docs/integrations/chat/google_generativeai',
    '[14] Prisma. (2024). "Prisma Documentation — ORM for Node.js and TypeScript." Available at: https://www.prisma.io/docs',
    '[15] Google Chrome Developers. (2024). "Chrome Extensions Manifest V3 Overview." Available at: https://developer.chrome.com/docs/extensions/mv3/intro/',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE   # spec: references single-spaced
    r = p.add_run(ref)
    set_font(r, 12)

# ════════════════════════════════════════════════════════════════════════════
# APPENDICES
# ════════════════════════════════════════════════════════════════════════════
appendix_heading("1", "SYSTEM INSTALLATION AND SETUP GUIDE")
body("This appendix provides step-by-step instructions for installing and configuring PhishGuardAI in a development environment.")
add_para("")
add_para("Prerequisites", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True)
bullet("Node.js 20 LTS or higher (https://nodejs.org)")
bullet("PostgreSQL 15 or higher (https://www.postgresql.org)")
bullet("Google Chrome 90+ browser")
bullet("Google AI Studio account for Gemini API key (https://aistudio.google.com)")
bullet("Git (https://git-scm.com)")

add_para("Backend Setup (Next.js API)", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True, space_before=8)
numbered("Clone the repository: git clone [repository-url] phishing-finder")
numbered("Navigate to the project directory: cd phishing-finder")
numbered("Install Node.js dependencies: npm install")
numbered("Create environment file: cp .env.example .env")
numbered("Configure .env variables: DATABASE_URL (PostgreSQL connection), GEMINI_API_KEY (from Google AI Studio), GEMINI_MODEL_NAME (gemini-2.0-flash), JWT_SECRET (random 32+ character string)")
numbered("Run Prisma migrations: npx prisma migrate dev")
numbered("Seed admin user: npm run seed:admin")
numbered("Start development server: npm run dev (runs on http://ec2-100-55-57-113.compute-1.amazonaws.com:3000)")

add_para("Chrome Extension Setup", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True, space_before=8)
numbered("Open Google Chrome and navigate to chrome://extensions")
numbered("Enable 'Developer mode' toggle (top right)")
numbered("Click 'Load unpacked' and select the chrome-extension directory")
numbered("The PhishGuardAI extension icon appears in the Chrome toolbar")
numbered("Click the extension icon and register/login with the API URL set to http://ec2-100-55-57-113.compute-1.amazonaws.com:3000")

add_para("Admin Dashboard Access", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True, space_before=8)
numbered("Navigate to http://ec2-100-55-57-113.compute-1.amazonaws.com:3000/admin/login in your browser")
numbered("Login with the admin credentials created by the seed script (admin@admin.com / Password@123)")
numbered("The URL Checks and Audit Records tabs are available upon successful login")

appendix_heading("2", "API RESPONSE EXAMPLES")
body("This appendix provides sample API request and response examples for the key endpoints.")
add_para("")

add_para("Sample: POST /api/auth/login", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True)
body('Request Body: { "email": "user@example.com", "password": "SecurePassword123" }')
body('Success Response (200): { "token": "eyJhbGciOiJIUzI1NiJ9...", "user": { "id": "uuid-here", "email": "user@example.com", "role": "USER" } }')
body('Error Response (401): { "error": "Invalid email or password" }')
add_para("")

add_para("Sample: POST /api/check-url", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True)
body('Request Body: { "url": "https://suspicious-site-example.com/login" }')
body('Headers: Authorization: Bearer eyJhbGciOiJIUzI1NiJ9...')
body('Success Response — Phishing (200): { "isPhishing": true, "hasWarning": false, "confidence": 0.95, "reason": "This site uses a domain name designed to impersonate a legitimate financial institution with a login form that submits credentials to an external server.", "warningType": [], "cached": false }')
body('Success Response — Warning (200): { "isPhishing": false, "hasWarning": true, "warningType": ["PIRACY"], "warningSeverity": "medium", "confidence": 0.82, "reason": "Site appears to host pirated software downloads.", "cached": false }')
body('Success Response — Safe (200): { "isPhishing": false, "hasWarning": false, "confidence": 0.98, "reason": "URL and content analysis confirm this is a legitimate, well-established website.", "cached": true }')
add_para("")

add_para("Sample: GET /api/admin/urls?page=1&limit=10&isPhishing=true", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True)
body('Headers: Authorization: Bearer [ADMIN-JWT-TOKEN]')
body('Success Response (200): { "data": [{ "id": "uuid", "url": "https://phishing-example.com", "isPhishing": true, "confidence": 0.96, "reason": "...", "checkedAt": "2024-01-15T10:30:00Z" }, ...], "total": 45, "page": 1, "limit": 10 }')

appendix_heading("3", "LANGCHAIN TOOL CALLING IMPLEMENTATION DETAILS")
body("This appendix provides detailed documentation of the LangChain tool calling implementation used in the PhishGuardAI analysis engine.")
add_para("")

add_para("Tool Definition (fetchWebsiteContent)", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True)
body("The fetchWebsiteContent tool is defined using LangChain's StructuredTool class with a Zod schema for input validation. The tool is registered with the Gemini model using the bindTools() method, enabling the model to autonomously decide when website inspection is needed.")
add_para("")

add_para("Analysis Loop Algorithm", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True)
numbered("Initialize LangChain ChatGoogleGenerativeAI with temperature=0.1 and tool binding")
numbered("Construct initial HumanMessage with URL and security analysis instructions")
numbered("Invoke model; receive AIMessage")
numbered("If AIMessage contains tool_calls: execute fetchWebsiteContent with provided URL argument")
numbered("Append ToolMessage with website content results to conversation")
numbered("Re-invoke model with updated message history")
numbered("Repeat steps 4-6 until no more tool calls (max 5 iterations)")
numbered("Parse final AIMessage content for JSON classification object")
numbered("Extract isPhishing, hasWarning, warningType, confidence, reason fields")
numbered("Store result in PostgreSQL UrlCheck table via Prisma upsert")

add_para("")
add_para("Content Extraction Patterns", WD_ALIGN_PARAGRAPH.LEFT, 12, bold=True)
bullet("Forms: <form[^>]*action=[\"']([^\"']*)[\"'][^>]*> — extracts form submission targets")
bullet("Scripts: <script[^>]*src=[\"']([^\"']*)[\"'] — extracts external script sources")
bullet("Inline JS: <script[^>]*>(.*?)</script> — extracts inline JavaScript for obfuscation analysis")
bullet("Links: <a[^>]*href=[\"']([^\"']*)[\"'] — extracts all anchor URLs for domain analysis")
bullet("Suspicious patterns: eval(, document.write(, atob(, String.fromCharCode( — common obfuscation indicators")

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
output_path = "/Users/karthikn/Dev/phishing-finder/PhishGuardAI_Project_Report.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
