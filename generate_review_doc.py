#!/usr/bin/env python3
"""Generate Second Review Document for AI-Based Phishing Detection System."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Second_Review_Document.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top, bottom, left, right with values like '1' (size in eighths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{val}" w:space="0" w:color="999999"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def style_header_row(row, bg_color="1a1a2e"):
    """Style a table header row with dark background and white bold text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(9)


def add_table_row(table, cells_text, bold_first=False):
    """Add a row to an existing table."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.bold = True
    return row


def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header styling."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(9)

    style_header_row(hdr)

    for row_data in rows:
        add_table_row(table, row_data)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if w and i < len(row.cells):
                    row.cells[i].width = Inches(w)

    # Reduce cell padding
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)

    return table


def add_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h


def add_body(doc, text):
    """Add a normal paragraph."""
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_screenshot_placeholder(doc, caption="Screenshot Placeholder"):
    """Add a grey box placeholder for screenshots."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {caption} ]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True
    # add a light border box via a 1-row table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.text = ""
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(f"\n\n[ {caption} ]\n\n")
    cr.font.size = Pt(12)
    cr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    cr.italic = True
    set_cell_shading(cell, "F5F5F5")
    for edge in ("top", "bottom", "left", "right"):
        set_cell_border(cell, **{edge: "4"})
    cell.width = Inches(5.5)
    doc.add_paragraph()  # spacer


def add_page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def build_title_page(doc):
    """Cover page."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Second Review Document · March 2026")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Based Phishing Detection System")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Chrome Browser Extension for Real-Time Phishing Protection")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(3):
        doc.add_paragraph()

    # Info table
    info = [
        ("Project Title", "AI-Based Phishing Detection System"),
        ("Review Stage", "Second Review"),
        ("Date", "March 2026"),
        ("Platform Name", "Phishing Finder"),
        ("Tech Stack", "Next.js · React · PostgreSQL · Prisma · Google Gemini AI · LangChain · Chrome Extension"),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info):
        c0 = tbl.cell(i, 0)
        c0.text = ""
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        c0.width = Inches(1.8)

        c1 = tbl.cell(i, 1)
        c1.text = ""
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.size = Pt(10)
        c1.width = Inches(4.5)

    add_page_break(doc)


def build_toc(doc):
    """Table of contents page."""
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Overall System Design",
        "2. Dataset (Database Schema)",
        "3. Modules Completed",
        "4. Pseudo Code for Incomplete Modules",
        "5. Test Cases",
        "6. Deviations and Justifications",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)
    add_page_break(doc)


def build_section1(doc):
    """Section 1: Overall System Design."""
    add_heading(doc, "1. Overall System Design", level=1)

    # 1.1 System Architecture
    add_heading(doc, "1.1 System Architecture", level=2)
    add_body(doc,
        "The AI-Based Phishing Detection System follows a multi-tier architecture consisting of "
        "four primary layers: a Chrome Extension frontend, a Next.js backend API server, a "
        "PostgreSQL database, and an external AI service layer powered by Google Gemini through "
        "the LangChain framework."
    )

    add_body(doc, "Architecture Layers")
    create_table(doc,
        ["Layer", "Technology", "Details"],
        [
            ["Client Layer", "Chrome Extension (Manifest V3)", "Background service worker · Content scripts · Popup UI · Scanning/Blocked/Warning pages"],
            ["API Layer", "Next.js 16 (Port 3000)", "Auth, Check-URL, Admin URLs, Admin Audit modules"],
            ["Database Layer", "PostgreSQL (Prisma 7)", "Primary relational data store for URL checks, users, audit records"],
            ["AI Service Layer", "Google Gemini 2.0 Flash + LangChain", "Phishing detection · Website content analysis · Confidence scoring"],
            ["External Services", "Google Gemini API", "AI-powered phishing analysis and warning categorization"],
        ],
        col_widths=[1.5, 2.0, 3.0],
    )

    doc.add_paragraph()

    # 1.2 Project Structure
    add_heading(doc, "1.2 Project Structure", level=2)
    add_body(doc,
        "The project is organized as a Next.js application with a separate Chrome Extension directory:"
    )

    structure_items = [
        ("chrome-extension/", " — Chrome Browser Extension (Manifest V3)"),
        ("  background.js", " — Service worker for URL interception and caching"),
        ("  popup.html/js", " — Extension popup with auth and scan controls"),
        ("  scanning.html/js", " — Interstitial page during AI analysis"),
        ("  blocked.html/js", " — Phishing block page"),
        ("  warning.html/js", " — Warning page for risky sites"),
        ("  content-warning.js", " — Warning banner injection script"),
        ("  options.html/js", " — Extension settings page"),
        ("src/app/", " — Next.js App Router"),
        ("  api/check-url/", " — Main phishing check API endpoint"),
        ("  api/auth/", " — Login and registration endpoints"),
        ("  api/admin/", " — Admin URL checks and audit record endpoints"),
        ("  admin/", " — Admin dashboard pages and components"),
        ("src/lib/", " — Core business logic"),
        ("  langchain-phishing.ts", " — LangChain + Gemini phishing detection engine"),
        ("  tools/fetch-website-content.ts", " — Website content analysis tool"),
        ("  auth.ts", " — JWT authentication utilities"),
        ("  middleware.ts", " — Auth middleware (requireAuth, requireAdmin)"),
        ("  prisma.ts", " — Prisma client singleton"),
        ("prisma/", " — Database schema and migrations"),
    ]
    for prefix, suffix in structure_items:
        p = doc.add_paragraph()
        run_b = p.add_run(prefix)
        run_b.bold = True
        run_b.font.size = Pt(9)
        run_b.font.name = "Courier New"
        run_s = p.add_run(suffix)
        run_s.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    doc.add_paragraph()

    # 1.3 Data Flow
    add_heading(doc, "1.3 Data Flow", level=2)

    add_body(doc, "URL Check Flow")
    add_body(doc,
        "User Navigates to URL → Extension Intercepts → Check In-Memory Cache → "
        "Check Database Cache (API) → If Not Cached: Redirect to Scanning Page → "
        "AI Analysis (LangChain + Gemini) → Tool Calling (Fetch Website Content) → "
        "Return Result → Cache in Database → Display Block Page / Warning Banner / Allow Navigation"
    )

    doc.add_paragraph()
    add_body(doc, "Admin Flow")
    add_body(doc,
        "Admin Login → View URL Checks Table (filter, sort, paginate) → "
        "View Audit Records (filter by action, user, date range) → "
        "Monitor Detection Statistics"
    )

    doc.add_paragraph()

    # 1.4 Technology Stack
    add_heading(doc, "1.4 Technology Stack", level=2)
    create_table(doc,
        ["Layer", "Technology", "Version", "Purpose"],
        [
            ["Frontend Framework", "Next.js", "16.0.4", "SSR/CSR React application + Admin panel"],
            ["UI Library", "React", "19.2.0", "Component rendering"],
            ["Styling", "Tailwind CSS", "v4", "Utility CSS"],
            ["UI Components", "shadcn/ui (Radix UI)", "Latest", "Accessible admin components"],
            ["Backend Framework", "Next.js API Routes", "16.0.4", "REST API server"],
            ["ORM", "Prisma", "7.2.0", "Type-safe database queries"],
            ["Database", "PostgreSQL", "15+", "Primary data store"],
            ["AI Engine", "Google Gemini", "2.0 Flash", "Phishing detection AI"],
            ["AI Framework", "LangChain", "0.3.x", "LLM orchestration + tool calling"],
            ["LangChain Gemini", "@langchain/google-genai", "0.0.20", "Gemini integration for LangChain"],
            ["Auth", "JWT + Bcrypt", "—", "Stateless authentication"],
            ["Extension", "Chrome Manifest V3", "—", "Browser extension platform"],
            ["Content Parsing", "Cheerio", "1.0.0", "HTML parsing for website analysis"],
            ["Validation", "Zod", "Latest", "Schema validation"],
            ["TypeScript", "TypeScript", "5.x", "Type-safe development"],
        ],
        col_widths=[1.5, 1.8, 0.8, 2.4],
    )

    doc.add_paragraph()

    # 1.5 API Design
    add_heading(doc, "1.5 API Design", level=2)
    add_bullet(doc, " http://ec2-100-55-57-113.compute-1.amazonaws.com:3000/api/", bold_prefix="Base URL:")
    add_bullet(doc, " JWT Bearer token required on protected routes", bold_prefix="Authentication:")
    add_bullet(doc, ' { data: T } wrapper with appropriate HTTP status codes', bold_prefix="Response Format:")
    add_bullet(doc, " Admin endpoints require ADMIN role in JWT payload", bold_prefix="Role Guard:")

    doc.add_paragraph()
    add_body(doc, "API Endpoints Summary")
    create_table(doc,
        ["Module", "Method", "Endpoint", "Access"],
        [
            ["Auth", "POST", "/api/auth/register", "Public"],
            ["Auth", "POST", "/api/auth/login", "Public"],
            ["URL Check", "POST", "/api/check-url", "Authenticated User"],
            ["Admin", "GET", "/api/admin/urls", "Admin Only"],
            ["Admin", "GET", "/api/admin/audit", "Admin Only"],
        ],
        col_widths=[1.2, 0.8, 2.5, 1.8],
    )

    doc.add_paragraph()

    add_body(doc, "POST /api/check-url Request Parameters")
    create_table(doc,
        ["Parameter", "Type", "Required", "Description"],
        [
            ["url", "String", "Yes", "The URL to analyze for phishing"],
            ["forceRefresh", "Boolean", "No", "Skip cache and force new AI analysis"],
            ["checkOnly", "Boolean", "No", "Only check cache, return notFound if not cached"],
        ],
        col_widths=[1.5, 1.0, 1.0, 3.0],
    )

    doc.add_paragraph()

    add_body(doc, "POST /api/check-url Response Fields")
    create_table(doc,
        ["Field", "Type", "Description"],
        [
            ["isPhishing", "Boolean", "Whether the URL is confirmed phishing"],
            ["hasWarning", "Boolean", "Whether the URL has security warnings"],
            ["warningType", "String?", "PIRACY / SCAMMING / RISKY_LINKS / SCAM_PRODUCTS / RISKY_FILES"],
            ["warningSeverity", "String?", "low / medium / high"],
            ["warningReason", "String?", "Detailed warning explanation"],
            ["reason", "String", "AI analysis reasoning"],
            ["confidence", "Float", "Confidence score (0.0 to 1.0)"],
            ["cached", "Boolean", "Whether result was from cache"],
        ],
        col_widths=[1.5, 1.0, 4.0],
    )

    add_page_break(doc)


def build_section2(doc):
    """Section 2: Dataset (Database Schema)."""
    add_heading(doc, "2. Dataset (Database Schema)", level=1)

    # 2.1 ER Overview
    add_heading(doc, "2.1 Entity Relationship Overview", level=2)
    add_body(doc, "The following entities form the core data model of the Phishing Detection System:")
    add_bullet(doc, " User → AuditRecord (one-to-many: one user has many audit records)")
    add_bullet(doc, " UrlCheck — standalone entity (shared cache, no foreign key relationships)")
    add_bullet(doc, " AuditRecord references User via userId foreign key with CASCADE delete")

    doc.add_paragraph()

    # 2.2 Database Models
    add_heading(doc, "2.2 Database Models", level=2)

    # User table
    add_body(doc, "Table: User")
    create_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, Default(uuid)", "Unique identifier"],
            ["email", "String", "Unique, Not Null", "Login email address"],
            ["passwordHash", "String", "Not Null", "Bcrypt hashed password"],
            ["role", "Enum", "USER / ADMIN", "User role for access control"],
            ["createdAt", "DateTime", "Default(now)", "Account creation timestamp"],
            ["updatedAt", "DateTime", "Auto-update", "Last update timestamp"],
        ],
        col_widths=[1.5, 1.0, 1.8, 2.2],
    )

    doc.add_paragraph()
    add_body(doc, "Relationships: One-to-many with AuditRecord")

    doc.add_paragraph()

    # UrlCheck table
    add_body(doc, "Table: UrlCheck")
    create_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, Default(uuid)", "Unique identifier"],
            ["url", "String", "Unique, Indexed, Not Null", "Normalized URL being checked"],
            ["isPhishing", "Boolean", "Not Null", "Whether URL is confirmed phishing"],
            ["hasWarning", "Boolean", "Default(false)", "Whether URL has security warnings"],
            ["warningType", "String?", "Nullable", "Warning category (PIRACY, SCAMMING, etc.)"],
            ["warningSeverity", "String?", "Nullable", "Severity level (low, medium, high)"],
            ["warningReason", "String?", "Nullable", "Detailed warning explanation"],
            ["confidence", "Float?", "Nullable, 0.0–1.0", "AI confidence score"],
            ["reason", "Text?", "Nullable", "Detailed AI analysis reasoning"],
            ["userId", "String?", "Nullable", "User who triggered the check"],
            ["checkedAt", "DateTime", "Default(now)", "When URL was analyzed"],
            ["createdAt", "DateTime", "Default(now)", "Record creation timestamp"],
            ["updatedAt", "DateTime", "Auto-update", "Record update timestamp"],
        ],
        col_widths=[1.5, 1.0, 1.8, 2.2],
    )

    doc.add_paragraph()
    add_body(doc, "Indexes: url (unique index for fast cache lookups)")

    doc.add_paragraph()

    # AuditRecord table
    add_body(doc, "Table: AuditRecord")
    create_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, Default(uuid)", "Unique identifier"],
            ["userId", "UUID", "FK → User.id, Not Null", "User who visited the URL"],
            ["url", "String", "Not Null", "URL that was visited"],
            ["isPhishing", "Boolean", "Not Null", "Whether URL was identified as phishing"],
            ["hasWarning", "Boolean", "Default(false)", "Whether URL had warnings"],
            ["warningType", "String?", "Nullable", "Warning category if applicable"],
            ["warningSeverity", "String?", "Nullable", "Severity level if applicable"],
            ["warningReason", "String?", "Nullable", "Warning explanation if applicable"],
            ["action", "Enum", "BLOCKED / WARNING_SHOWN / ALLOWED", "Action taken by the system"],
            ["ipAddress", "String?", "Nullable", "IP address of the user"],
            ["userAgent", "String?", "Nullable", "Browser user agent string"],
            ["visitedAt", "DateTime", "Default(now)", "Timestamp of the visit"],
            ["createdAt", "DateTime", "Default(now)", "Record creation timestamp"],
        ],
        col_widths=[1.5, 1.0, 1.8, 2.2],
    )

    doc.add_paragraph()
    add_body(doc, "Indexes: userId, visitedAt, action")
    add_body(doc, "Constraint: CASCADE delete — when a user is deleted, all audit records are removed")

    doc.add_paragraph()

    # 2.3 Sample Data
    add_heading(doc, "2.3 Sample Data", level=2)

    add_body(doc, "Sample Users")
    create_table(doc,
        ["ID", "Email", "Role"],
        [
            ["1", "admin@admin.com", "ADMIN"],
            ["2", "user@example.com", "USER"],
            ["3", "alice@example.com", "USER"],
        ],
    )

    doc.add_paragraph()

    add_body(doc, "Sample URL Checks")
    create_table(doc,
        ["ID", "URL", "Is Phishing", "Has Warning", "Warning Type", "Confidence"],
        [
            ["1", "https://paypal-secure-login.xyz", "true", "false", "—", "0.95"],
            ["2", "https://free-movies-stream.com", "false", "true", "PIRACY", "0.88"],
            ["3", "https://google.com", "false", "false", "—", "0.99"],
        ],
    )

    doc.add_paragraph()

    add_body(doc, "Sample Audit Records")
    create_table(doc,
        ["ID", "User Email", "URL", "Action", "Is Phishing"],
        [
            ["1", "user@example.com", "https://paypal-secure-login.xyz", "BLOCKED", "true"],
            ["2", "user@example.com", "https://free-movies-stream.com", "WARNING_SHOWN", "false"],
            ["3", "alice@example.com", "https://google.com", "ALLOWED", "false"],
        ],
    )

    add_page_break(doc)


def build_section3(doc):
    """Section 3: Modules Completed."""
    add_heading(doc, "3. Modules Completed", level=1)

    # Module 1: Authentication System
    add_heading(doc, "Module 1: Authentication System", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: COMPLETE")
    r.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)

    add_body(doc,
        "Full JWT-based authentication with role separation (USER and ADMIN). "
        "Passwords hashed with Bcrypt. Token stored in Chrome Extension local storage "
        "and admin panel localStorage."
    )

    add_body(doc, "Features:")
    add_bullet(doc, "User registration with email and password")
    add_bullet(doc, "Login returns JWT access token (7-day expiry) with userId, email, and role in payload")
    add_bullet(doc, "Role-based access control: USER role for extension users, ADMIN role for dashboard")
    add_bullet(doc, "requireAuth middleware validates JWT on all protected API routes")
    add_bullet(doc, "requireAdmin middleware enforces ADMIN role on admin endpoints")
    add_bullet(doc, "Extension stores auth token in chrome.storage and sends as Bearer header")
    add_bullet(doc, "Admin panel stores token in localStorage with admin-specific key")
    add_bullet(doc, "Fail-open policy: extension allows navigation if user is not logged in")

    add_body(doc, "Screenshots:")
    add_screenshot_placeholder(doc, "Extension Login/Register Popup")
    add_screenshot_placeholder(doc, "Admin Login Page")

    # Module 2: Chrome Extension
    add_heading(doc, "Module 2: Chrome Browser Extension", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: COMPLETE")
    r.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)

    add_body(doc,
        "Fully functional Chrome browser extension built with Manifest V3. The extension "
        "automatically intercepts all web navigation, performs real-time phishing detection, "
        "and provides multiple levels of user protection."
    )

    add_body(doc, "Features:")
    add_bullet(doc, "Manifest V3 with permissions: webRequest, tabs, storage, activeTab")
    add_bullet(doc, "Background service worker intercepts all tab navigation via chrome.tabs.onUpdated")
    add_bullet(doc, "In-memory URL cache for instant responses on previously checked domains")
    add_bullet(doc, "Popup interface with login/register, current page status, and manual scan button")
    add_bullet(doc, "Scanning interstitial page shown during AI analysis with loading animation")
    add_bullet(doc, "Blocked page for confirmed phishing sites with detailed reason and options (Scan Again, Go Back, Proceed Anyway)")
    add_bullet(doc, "Warning page for risky sites with warning type, severity, and reason")
    add_bullet(doc, "Content script injects warning banner overlay on risky websites")
    add_bullet(doc, "Options page with toggles: Auto Scan (on/off), Block Sites with Warnings (on/off)")
    add_bullet(doc, "Settings persisted to chrome.storage.sync")
    add_bullet(doc, "Skips internal Chrome URLs, extension pages, and localhost")

    add_body(doc, "Screenshots:")
    add_screenshot_placeholder(doc, "Extension Popup — Logged In with Scan Result")
    add_screenshot_placeholder(doc, "Scanning Interstitial Page")
    add_screenshot_placeholder(doc, "Blocked Page — Phishing Site Detected")
    add_screenshot_placeholder(doc, "Warning Banner on Risky Website")
    add_screenshot_placeholder(doc, "Extension Options Page")

    # Module 3: URL Check Engine
    add_heading(doc, "Module 3: URL Check Engine", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: COMPLETE")
    r.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)

    add_body(doc,
        "Core phishing detection API with multi-layer caching and AI-powered analysis. "
        "The engine normalizes URLs to base domains, checks a two-tier cache (in-memory + database), "
        "and invokes AI analysis for new URLs."
    )

    add_body(doc, "Features:")
    add_bullet(doc, "POST /api/check-url endpoint with JWT authentication")
    add_bullet(doc, "URL normalization: lowercase, strip www prefix, extract base domain")
    add_bullet(doc, "Two-tier caching: extension in-memory cache + PostgreSQL database cache")
    add_bullet(doc, "forceRefresh parameter to bypass cache and trigger fresh AI analysis")
    add_bullet(doc, "checkOnly parameter for cache-only lookups (returns notFound if not cached)")
    add_bullet(doc, "Upsert UrlCheck record after each analysis")
    add_bullet(doc, "Automatic AuditRecord creation for BLOCKED and WARNING_SHOWN actions")
    add_bullet(doc, "Fail-open policy: allows navigation if API check fails")

    add_body(doc, "Screenshots:")
    add_screenshot_placeholder(doc, "URL Check API Response — Phishing Detected")
    add_screenshot_placeholder(doc, "URL Check API Response — Safe Site")

    # Module 4: AI Integration
    add_heading(doc, "Module 4: AI Integration (LangChain + Google Gemini)", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: COMPLETE")
    r.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)

    add_body(doc,
        "Advanced AI-powered phishing detection using LangChain framework with Google Gemini 2.0 Flash model. "
        "The system uses tool calling to fetch and analyze actual website content when initial URL analysis "
        "yields low confidence."
    )

    add_body(doc, "Features:")
    add_bullet(doc, "LangChain ChatGoogleGenerativeAI model with temperature 0.1 for consistent results")
    add_bullet(doc, "Sophisticated system prompt instructing AI to act as expert cybersecurity analyst")
    add_bullet(doc, "Tool calling mechanism: FetchWebsiteContentTool for deep website analysis")
    add_bullet(doc, "Iterative analysis loop (up to 5 iterations) for complex cases")
    add_bullet(doc, "Website content analysis: HTML structure, JavaScript code, forms, links, meta tags")
    add_bullet(doc, "Suspicious pattern detection: eval(), base64 encoding, obfuscated code, hidden elements")
    add_bullet(doc, "Form analysis: password fields, hidden forms, suspicious action URLs")
    add_bullet(doc, "Link analysis: internal vs external links, typosquatting detection")
    add_bullet(doc, "Confidence scoring from 0.0 to 1.0 — triggers tool call if below 0.8 threshold")
    add_bullet(doc, "JSON response parsing with fallback text extraction")
    add_bullet(doc, "Token limit management (~10k tokens) with smart content trimming")

    add_body(doc, "Screenshots:")
    add_screenshot_placeholder(doc, "AI Analysis Result with High Confidence")
    add_screenshot_placeholder(doc, "AI Analysis with Tool Calling — Website Content Fetched")

    # Module 5: Warning Categorization System
    add_heading(doc, "Module 5: Warning Categorization System", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: COMPLETE")
    r.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)

    add_body(doc,
        "Multi-category warning system that classifies risky websites beyond simple phishing detection. "
        "Provides nuanced security guidance through five distinct warning categories with severity levels."
    )

    add_body(doc, "Features:")
    add_bullet(doc, "Five warning categories:")
    add_bullet(doc, "  PIRACY — Sites hosting pirated content or illegal downloads")
    add_bullet(doc, "  SCAMMING — Fraudulent services, fake reviews, deceptive practices")
    add_bullet(doc, "  RISKY_LINKS — Suspicious redirects, link farms, suspicious external links")
    add_bullet(doc, "  SCAM_PRODUCTS — Fake products, counterfeit goods, scam marketplaces")
    add_bullet(doc, "  RISKY_FILES — Potentially malicious file distribution sites")
    add_bullet(doc, "Three severity levels: low, medium, high")
    add_bullet(doc, "Warning banner injected into page via content script with type and severity display")
    add_bullet(doc, "Dedicated warning page with full details and options to proceed or go back")
    add_bullet(doc, "Optional: Block sites with warnings setting in extension options")

    add_body(doc, "Screenshots:")
    add_screenshot_placeholder(doc, "Warning Banner — Piracy Site (Medium Severity)")
    add_screenshot_placeholder(doc, "Warning Page — Scam Products (High Severity)")

    # Module 6: Admin Dashboard
    add_heading(doc, "Module 6: Admin Dashboard", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: COMPLETE")
    r.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)

    add_body(doc,
        "Comprehensive web-based admin dashboard for monitoring all phishing detection activity. "
        "Built with Next.js, shadcn/ui components, and Tailwind CSS with a tabbed interface."
    )

    add_body(doc, "Features:")
    add_bullet(doc, "Admin login with separate authentication (ADMIN role required)")
    add_bullet(doc, "URL Checks tab: paginated table of all analyzed URLs")
    add_bullet(doc, "  Columns: URL, Phishing status, Warning status, Warning Type, Confidence, Checked At")
    add_bullet(doc, "  Filters: search by URL, filter by phishing status, filter by warning status")
    add_bullet(doc, "  Sortable columns (ascending/descending)")
    add_bullet(doc, "  Configurable page sizes: 25, 50, or 100 records per page")
    add_bullet(doc, "Audit Records tab: paginated table of all user interactions")
    add_bullet(doc, "  Columns: User Email, URL, Action, Phishing, Warning, Warning Details, IP, Visited At")
    add_bullet(doc, "  Filters: user search, action type (BLOCKED/WARNING_SHOWN/ALLOWED), phishing status, date range")
    add_bullet(doc, "Responsive design with modern UI (Radix UI, Tailwind CSS)")
    add_bullet(doc, "Header with user info and logout button")

    add_body(doc, "Screenshots:")
    add_screenshot_placeholder(doc, "Admin Dashboard — URL Checks Tab")
    add_screenshot_placeholder(doc, "Admin Dashboard — Audit Records Tab")
    add_screenshot_placeholder(doc, "Admin Dashboard — Filters Applied")

    # Module 7: Audit Logging
    add_heading(doc, "Module 7: Audit Logging System", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: COMPLETE")
    r.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)

    add_body(doc,
        "Automatic audit trail for all URL check interactions. Every time a user visits a URL "
        "that results in a block or warning action, a detailed audit record is created."
    )

    add_body(doc, "Features:")
    add_bullet(doc, "Automatic audit record creation on BLOCKED and WARNING_SHOWN actions")
    add_bullet(doc, "Records user identification (userId), visited URL, action taken")
    add_bullet(doc, "Captures phishing status, warning details (type, severity, reason)")
    add_bullet(doc, "Stores IP address and browser user agent for forensic analysis")
    add_bullet(doc, "Timestamped with visitedAt for chronological tracking")
    add_bullet(doc, "Cascade delete: records removed when associated user is deleted")
    add_bullet(doc, "Indexed fields (userId, visitedAt, action) for efficient querying")

    add_body(doc, "Screenshots:")
    add_screenshot_placeholder(doc, "Audit Records in Admin Dashboard")

    add_page_break(doc)


def build_section4(doc):
    """Section 4: Pseudo Code for Incomplete Modules."""
    add_heading(doc, "4. Pseudo Code for Incomplete Modules", level=1)

    # 4.1 Rate Limiting
    add_heading(doc, "4.1 API Rate Limiting", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: Not Implemented")
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

    add_body(doc,
        "Rate limiting to prevent abuse and protect against denial-of-service attacks. "
        "Includes per-user request limits, IP-based throttling, and AI API call management."
    )

    pseudo = (
        "FUNCTION rateLimitMiddleware(request):\n"
        "    userId = extractUserId(request)\n"
        "    ipAddress = extractIP(request)\n"
        "    endpoint = request.path\n"
        "\n"
        "    // Per-user rate limit: 100 requests per minute\n"
        "    userKey = 'rate:user:' + userId + ':' + currentMinute()\n"
        "    userCount = redis.INCR(userKey)\n"
        "    IF userCount == 1: redis.EXPIRE(userKey, 60)\n"
        "    IF userCount > 100: RETURN 429 Too Many Requests\n"
        "\n"
        "    // IP-based rate limit: 200 requests per minute\n"
        "    ipKey = 'rate:ip:' + ipAddress + ':' + currentMinute()\n"
        "    ipCount = redis.INCR(ipKey)\n"
        "    IF ipCount == 1: redis.EXPIRE(ipKey, 60)\n"
        "    IF ipCount > 200: RETURN 429 Too Many Requests\n"
        "\n"
        "    // AI API rate limit: 30 AI calls per user per hour\n"
        "    IF endpoint == '/api/check-url' AND NOT checkOnly:\n"
        "        aiKey = 'rate:ai:' + userId + ':' + currentHour()\n"
        "        aiCount = redis.INCR(aiKey)\n"
        "        IF aiCount == 1: redis.EXPIRE(aiKey, 3600)\n"
        "        IF aiCount > 30: RETURN 429 AI Rate Limit Exceeded\n"
        "\n"
        "    NEXT()\n"
        "END FUNCTION"
    )
    p = doc.add_paragraph()
    run = p.add_run(pseudo)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    doc.add_paragraph()

    # 4.2 Token Refresh
    add_heading(doc, "4.2 Token Refresh Mechanism", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: Not Implemented")
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

    add_body(doc,
        "JWT token refresh to maintain user sessions without frequent re-authentication. "
        "Uses short-lived access tokens with long-lived refresh tokens for improved security."
    )

    pseudo = (
        "FUNCTION refreshToken(request):\n"
        "    refreshToken = request.body.refreshToken\n"
        "    payload = verifyRefreshToken(refreshToken)\n"
        "    IF payload IS NULL: RETURN 401 Invalid Refresh Token\n"
        "\n"
        "    // Check if refresh token is in whitelist\n"
        "    isValid = db.RefreshToken.findUnique({ token: refreshToken })\n"
        "    IF isValid IS NULL: RETURN 401 Token Revoked\n"
        "\n"
        "    // Rotate: delete old, issue new pair\n"
        "    db.RefreshToken.delete({ token: refreshToken })\n"
        "    newAccessToken = generateAccessToken(payload.userId, '15m')\n"
        "    newRefreshToken = generateRefreshToken(payload.userId, '30d')\n"
        "    db.RefreshToken.create({ userId: payload.userId, token: newRefreshToken })\n"
        "\n"
        "    RETURN { accessToken: newAccessToken, refreshToken: newRefreshToken }\n"
        "END FUNCTION"
    )
    p = doc.add_paragraph()
    run = p.add_run(pseudo)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    doc.add_paragraph()

    # 4.3 Advanced Analytics
    add_heading(doc, "4.3 Advanced Analytics Dashboard", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: Not Implemented")
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

    add_body(doc,
        "Enhanced admin dashboard with visual analytics showing detection trends, "
        "threat distribution, and system performance metrics."
    )

    pseudo = (
        "FUNCTION getAnalytics(timeRange):\n"
        "    analytics = {\n"
        "        totalUrlsChecked: COUNT(UrlCheck),\n"
        "        totalPhishingDetected: COUNT(UrlCheck WHERE isPhishing = true),\n"
        "        totalWarningsIssued: COUNT(UrlCheck WHERE hasWarning = true),\n"
        "        totalSafeSites: COUNT(UrlCheck WHERE isPhishing = false AND hasWarning = false),\n"
        "        averageConfidence: AVG(UrlCheck.confidence),\n"
        "    }\n"
        "\n"
        "    // Detection trend (daily counts for last 30 days)\n"
        "    FOR EACH day IN last30Days:\n"
        "        analytics.trend.append({\n"
        "            date: day,\n"
        "            phishing: COUNT(UrlCheck WHERE date = day AND isPhishing),\n"
        "            warnings: COUNT(UrlCheck WHERE date = day AND hasWarning),\n"
        "            safe: COUNT(UrlCheck WHERE date = day AND NOT isPhishing AND NOT hasWarning)\n"
        "        })\n"
        "\n"
        "    // Warning type distribution (pie chart data)\n"
        "    analytics.warningDistribution = GROUP BY warningType, COUNT(*)\n"
        "\n"
        "    // Top phishing domains\n"
        "    analytics.topPhishingDomains = TOP 10 UrlCheck WHERE isPhishing ORDER BY checkedAt DESC\n"
        "\n"
        "    // Active users (last 7 days)\n"
        "    analytics.activeUsers = COUNT(DISTINCT AuditRecord.userId WHERE visitedAt > 7daysAgo)\n"
        "\n"
        "    RETURN analytics\n"
        "END FUNCTION"
    )
    p = doc.add_paragraph()
    run = p.add_run(pseudo)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    doc.add_paragraph()

    # 4.4 Multi-Browser Support
    add_heading(doc, "4.4 Multi-Browser Extension Support", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Status: Not Implemented")
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

    add_body(doc,
        "Extend the browser extension to support Firefox, Edge, and Safari browsers "
        "using the WebExtensions API standard."
    )

    pseudo = (
        "FUNCTION buildForBrowser(targetBrowser):\n"
        "    // Shared core logic\n"
        "    coreModules = ['api-client.js', 'auth.js', 'content-warning.js']\n"
        "\n"
        "    IF targetBrowser == 'firefox':\n"
        "        manifest = generateManifestV2({\n"
        "            background: { scripts: ['background.js'] },\n"
        "            permissions: ['webRequest', 'tabs', 'storage', '<all_urls>'],\n"
        "            browser_specific_settings: { gecko: { id: 'phishing-finder@ext' } }\n"
        "        })\n"
        "        replaceChromAPIs(coreModules, 'browser.*')\n"
        "\n"
        "    ELSE IF targetBrowser == 'edge':\n"
        "        manifest = generateManifestV3()  // Same as Chrome\n"
        "        // Edge supports Chrome extension APIs natively\n"
        "\n"
        "    ELSE IF targetBrowser == 'safari':\n"
        "        manifest = generateSafariManifest()\n"
        "        wrapWithSafariWebExtensionBridge(coreModules)\n"
        "\n"
        "    outputDir = 'dist/' + targetBrowser + '/'\n"
        "    copyFiles(coreModules + htmlPages + cssFiles, outputDir)\n"
        "    writeFile(outputDir + 'manifest.json', manifest)\n"
        "    package(outputDir)\n"
        "END FUNCTION"
    )
    p = doc.add_paragraph()
    run = p.add_run(pseudo)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    add_page_break(doc)


def build_section5(doc):
    """Section 5: Test Cases."""
    add_heading(doc, "5. Test Cases", level=1)

    # 5.1 Authentication
    add_heading(doc, "5.1 Authentication Module", level=2)
    create_table(doc,
        ["TC#", "Test Case", "Input", "Expected Output", "Status"],
        [
            ["TC-01", "Register with valid data", "email, password", "201 Created, user object + JWT returned", "PASS"],
            ["TC-02", "Register with duplicate email", "Existing email", "409 Conflict", "PASS"],
            ["TC-03", "Login with valid credentials", "email, password", "200 OK, access_token returned", "PASS"],
            ["TC-04", "Login with wrong password", "email, wrong password", "401 Unauthorized", "PASS"],
            ["TC-05", "Access protected route without token", "No Authorization header", "401 Unauthorized", "PASS"],
            ["TC-06", "Access protected route with valid token", "Bearer token in header", "200 OK, data returned", "PASS"],
            ["TC-07", "Admin login with USER role", "USER role credentials", "403 Forbidden (not admin)", "PASS"],
            ["TC-08", "Admin login with ADMIN role", "ADMIN role credentials", "200 OK, admin access granted", "PASS"],
            ["TC-09", "Access admin route with USER token", "USER JWT", "403 Forbidden", "PASS"],
        ],
        col_widths=[0.5, 1.8, 1.5, 1.8, 0.6],
    )

    doc.add_paragraph()

    # 5.2 URL Check Module
    add_heading(doc, "5.2 URL Check Module", level=2)
    create_table(doc,
        ["TC#", "Test Case", "Input", "Expected Output", "Status"],
        [
            ["TC-10", "Check known phishing URL", "POST /api/check-url, phishing URL", "isPhishing: true, confidence > 0.8", "PASS"],
            ["TC-11", "Check safe URL (e.g., google.com)", "POST /api/check-url, google.com", "isPhishing: false, hasWarning: false", "PASS"],
            ["TC-12", "Check URL with warning", "POST /api/check-url, piracy site", "hasWarning: true, warningType set", "PASS"],
            ["TC-13", "Check cached URL (no AI call)", "Same URL as TC-10", "cached: true, instant response", "PASS"],
            ["TC-14", "Force refresh bypasses cache", "forceRefresh: true", "cached: false, fresh AI analysis", "PASS"],
            ["TC-15", "checkOnly returns notFound", "checkOnly: true, new URL", "{ notFound: true }", "PASS"],
            ["TC-16", "checkOnly returns cached result", "checkOnly: true, known URL", "Full result with cached: true", "PASS"],
            ["TC-17", "URL normalization consistency", "http://WWW.Example.COM/path", "Normalized to https://example.com", "PASS"],
            ["TC-18", "Check URL without auth token", "No Bearer token", "401 Unauthorized", "PASS"],
            ["TC-19", "Invalid URL format", "Not a valid URL", "400 Bad Request", "PASS"],
        ],
        col_widths=[0.5, 1.8, 1.5, 1.8, 0.6],
    )

    doc.add_paragraph()

    # 5.3 Chrome Extension
    add_heading(doc, "5.3 Chrome Extension Module", level=2)
    create_table(doc,
        ["TC#", "Test Case", "Input", "Expected Output", "Status"],
        [
            ["TC-20", "Extension intercepts navigation", "Navigate to any URL", "Background worker triggers check", "PASS"],
            ["TC-21", "Phishing site blocked", "Navigate to phishing URL", "Redirected to blocked.html", "PASS"],
            ["TC-22", "Warning banner shown", "Navigate to risky URL", "Warning banner injected into page", "PASS"],
            ["TC-23", "Safe site allowed", "Navigate to safe URL", "Normal navigation proceeds", "PASS"],
            ["TC-24", "Scanning page displayed", "Navigate to uncached URL", "scanning.html shown during analysis", "PASS"],
            ["TC-25", "In-memory cache works", "Revisit checked URL", "Instant response, no API call", "PASS"],
            ["TC-26", "Popup shows scan result", "Click extension icon", "Current page status displayed", "PASS"],
            ["TC-27", "Manual scan button works", "Click 'Scan Current Page'", "Fresh scan triggered with result", "PASS"],
            ["TC-28", "Auto-scan toggle off", "Disable auto-scan in options", "No automatic interception", "PASS"],
            ["TC-29", "Proceed anyway from blocked page", "Click 'Proceed Anyway'", "Navigates to blocked URL after confirmation", "PASS"],
            ["TC-30", "Extension skip chrome:// URLs", "Navigate to chrome://settings", "No interception, normal navigation", "PASS"],
        ],
        col_widths=[0.5, 1.8, 1.5, 1.8, 0.6],
    )

    doc.add_paragraph()

    # 5.4 AI Integration
    add_heading(doc, "5.4 AI Integration Module", level=2)
    create_table(doc,
        ["TC#", "Test Case", "Input", "Expected Output", "Status"],
        [
            ["TC-31", "AI detects phishing URL", "Known phishing domain", "isPhishing: true, confidence > 0.8", "PASS"],
            ["TC-32", "AI detects safe URL", "Legitimate domain (google.com)", "isPhishing: false, confidence > 0.9", "PASS"],
            ["TC-33", "Tool calling triggered", "Ambiguous URL (confidence < 0.8)", "FetchWebsiteContent tool invoked", "PASS"],
            ["TC-34", "Website content analyzed", "URL with suspicious forms", "Suspicious indicators detected", "PASS"],
            ["TC-35", "Warning type categorized", "Piracy streaming site", "warningType: PIRACY", "PASS"],
            ["TC-36", "Severity assigned correctly", "High-risk scam site", "warningSeverity: high", "PASS"],
            ["TC-37", "Confidence score returned", "Any URL", "Float between 0.0 and 1.0", "PASS"],
            ["TC-38", "AI provides reasoning", "Any URL", "Non-empty reason field", "PASS"],
            ["TC-39", "Iterative analysis (multi-turn)", "Complex URL requiring tool call", "Multiple iterations before final result", "PASS"],
            ["TC-40", "Token limit management", "URL with large HTML content", "Content trimmed to ~10k tokens", "PASS"],
        ],
        col_widths=[0.5, 1.8, 1.5, 1.8, 0.6],
    )

    doc.add_paragraph()

    # 5.5 Admin Module
    add_heading(doc, "5.5 Admin Dashboard Module", level=2)
    create_table(doc,
        ["TC#", "Test Case", "Input", "Expected Output", "Status"],
        [
            ["TC-41", "Admin views URL checks", "GET /api/admin/urls", "Paginated URL check list", "PASS"],
            ["TC-42", "Admin filters by phishing", "isPhishing=true query param", "Only phishing URLs shown", "PASS"],
            ["TC-43", "Admin filters by warning", "hasWarning=true query param", "Only warned URLs shown", "PASS"],
            ["TC-44", "Admin searches URLs", "search=paypal query param", "URLs containing 'paypal'", "PASS"],
            ["TC-45", "Admin sorts by confidence", "sortBy=confidence, sortOrder=desc", "URLs sorted by confidence descending", "PASS"],
            ["TC-46", "Admin views audit records", "GET /api/admin/audit", "Paginated audit record list", "PASS"],
            ["TC-47", "Admin filters by action", "action=BLOCKED", "Only BLOCKED records shown", "PASS"],
            ["TC-48", "Admin filters by date range", "dateFrom, dateTo params", "Records within date range", "PASS"],
            ["TC-49", "Admin searches by user", "userSearch=user@example.com", "Records for that user only", "PASS"],
            ["TC-50", "Pagination works correctly", "page=2, limit=25", "Second page with 25 records", "PASS"],
            ["TC-51", "Non-admin access denied", "USER role JWT on admin route", "403 Forbidden", "PASS"],
        ],
        col_widths=[0.5, 1.8, 1.5, 1.8, 0.6],
    )

    add_page_break(doc)


def build_section6(doc):
    """Section 6: Deviations and Justifications."""
    add_heading(doc, "6. Deviations and Justifications", level=1)

    create_table(doc,
        ["#", "Original Plan", "Deviation", "Justification"],
        [
            [
                "D-01",
                "OpenAI GPT for AI-powered phishing detection",
                "Changed to Google Gemini 2.0 Flash",
                "Gemini provides a generous free-tier API with no billing required during development. "
                "Gemini 2.0 Flash offers fast response times ideal for real-time browser extension use. "
                "The LangChain abstraction layer allows easy model swapping if needed."
            ],
            [
                "D-02",
                "Multi-browser support (Chrome, Firefox, Edge, Safari)",
                "Chrome-only extension",
                "Scope limited to Chrome for the project review. Chrome holds the largest market share "
                "and Manifest V3 is the most mature extension platform. The backend API is browser-agnostic, "
                "so extensions for other browsers can be added with minimal backend changes."
            ],
            [
                "D-03",
                "Cloud-hosted deployment (AWS/GCP)",
                "Local development environment only",
                "Cloud deployment requires paid subscriptions and infrastructure management that is "
                "outside the academic scope. The application runs identically in local and cloud "
                "environments via standard Next.js deployment. Docker configuration is prepared for "
                "future containerized deployment."
            ],
            [
                "D-04",
                "Email phishing detection (email content analysis)",
                "Feature removed — web-only phishing detection",
                "Email phishing detection requires IMAP/SMTP integration and email provider access "
                "which adds significant complexity. The project focuses on web-based phishing detection "
                "through browser extension, which addresses the most common phishing attack vector."
            ],
            [
                "D-05",
                "Real-time analytics dashboard with charts and graphs",
                "Basic tabular admin dashboard implemented",
                "The core admin functionality (URL checks, audit records, filtering, sorting, pagination) "
                "is fully implemented. Visual analytics (charts, trend graphs) are planned for a future "
                "iteration. Pseudo code is provided in Section 4."
            ],
            [
                "D-06",
                "Rate limiting and advanced security features",
                "Not implemented",
                "Basic authentication and authorization are fully implemented with JWT + role-based access. "
                "Rate limiting requires Redis infrastructure. Security hardening (CSRF, CSP headers, "
                "rate limiting) is planned for the production deployment phase. Pseudo code is provided "
                "in Section 4."
            ],
            [
                "D-07",
                "Mobile application for phishing detection",
                "Not implemented",
                "Scope limited to Chrome browser extension for the academic project. The Next.js admin "
                "panel is responsive and mobile-accessible via browser. A dedicated mobile app was never "
                "part of the core requirement."
            ],
        ],
        col_widths=[0.5, 1.5, 1.5, 3.0],
    )


# ---------------------------------------------------------------------------
# Header / Footer
# ---------------------------------------------------------------------------

def add_header_footer(doc):
    """Add header and footer to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = "Second Review Document · March 2026"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = ""
        run = fp.add_run("AI-Based Phishing Detection System — Second Review")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Base styles
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # Adjust heading styles
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Build document
    build_title_page(doc)
    build_toc(doc)
    build_section1(doc)
    build_section2(doc)
    build_section3(doc)
    build_section4(doc)
    build_section5(doc)
    build_section6(doc)
    add_header_footer(doc)

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
